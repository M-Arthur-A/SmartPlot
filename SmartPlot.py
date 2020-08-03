# -*- coding: utf-8 -*-
### qpy:kivy
# pip install https://github.com/pyinstaller/pyinstaller/archive/develop.zip
# C:\Users\1\PycharmProjects\SmartPlot\venv\Scripts\python.exe C:\Anaconda3\Scripts\pyinstaller.exe SmartPlot.spec

import pandas as pd
import difflib
import re
from glob import glob as gb
from lxml import etree
import sys
from kivy.config import Config
from kivy.app import App
from kivy.lang import Builder
from kivy.uix.button import Button
from kivy.uix.checkbox import CheckBox
from kivy.core.window import Window
from kivy.uix.image import Image
from kivy.uix.popup import Popup
from kivy.uix.dropdown import DropDown
from kivy.uix.screenmanager import ScreenManager, Screen, FadeTransition
from kivy.graphics import RoundedRectangle
from kivy.uix.scrollview import ScrollView
from kivy.uix.tabbedpanel import TabbedPanel
from kivy.properties import ListProperty, StringProperty
from kivy.uix.video import Video
from kivy.clock import Clock
import datetime as dt
from kivy.metrics import dp


sys.setrecursionlimit(10 ** 9)
Config.set('graphics', 'resizable', False)


main_widget_kv = '''
TabbedPanel:
    id: thePanel
    tab_height: 20
    do_default_tab: False
    # background_color: (0,0,0,1)
    canvas:
        Color:
            rgb: (.1,.1,.1,1)
        Rectangle:
            size: self.size
    TabbedPanelItem:
        id: tab1
        text:'SmartPlot'
        ScreenManager:
            id: scr_mngr
            Screen:
                name: 'SP'
                BoxLayout:
                    id: main
                    orientation: 'vertical'
                    canvas.before:
                        Rectangle:
                            size: self.size
                            source: 'Emina/Enable/Default'
                    BoxLayout:
                        orientation: 'vertical'
                        size_hint: None, None
                        size: (800, 120)
                        canvas.before:
                            Rectangle:
                                pos: (15, 489)
                                size: (324, 81)
                                source: 'Emina/Enable/Default_1'
                        BoxLayout:
                            Label:
                                text: "   SMARTPLOT"
                                size_hint_x: 11.8
                                pos_hint: {"x":1, "y":-0.3}
                                text_size: self.size
                                halign: 'left'
                                valign: 'bottom'
                                font_size: 48
                                color: 0.3,0.3,0.3,1
                            BoxLayout:
                                size_hint: None, None
                                size: dp(85), dp(85)
                                pos_hint:{'top': 1.06, 'right': 1}
                                RoundedButton:
                                    background_normal: 'Emina/Enable/Help'
                                    background_down: 'Emina/Activated/Help'
                                    halign: 'right'
                                    text: "?"
                                    font_size: 50
                                    pos: (1,0.5)
                                    on_release:
                                        scr_mngr.transition.direction = 'left'
                                        scr_mngr.current = 'Info'
                        Label:
                            text: "Активы по группам      "
                            text_size: self.size
                            halign: 'right'
                            color: 0.3,0.3,0.3,1
                            font_size: 32
                            canvas.before:
                                Rectangle:
                                    pos: (427, 448)
                                    size: (367, 56)
                                    source: 'Emina/Enable/Default_2'
                    BoxLayout:
                        # size_hint_y: 1.45
                        on_dropfile:
                        orientation: 'vertical'
                        canvas.before:
                            Rectangle:
                                pos: (290, 215)
                                size: (217, 223)
                                source: 'Emina/Enable/Default_3'
                        BoxLayout:
                        Label:
                            size_hint_y: 0.2
                            text: "Drag'n Drop"
                            color: 0.3,0.3,0.3,1
                        Image:
                            id: doc_img
                            source: 'Emina/Enable/DnD'
                        BoxLayout:
                            size_hint_y: 0.8
                    BoxLayout:
                        #size_hint_y: 0.5
                        size_hint: None, None
                        size: (800, 93)
                        padding: [40,40,40,0]
                        canvas.before:
                            Rectangle:
                                pos: (0, 115)
                                size: (800, 94)
                                source: 'Emina/Enable/Default_4'
                        BoxLayout:
                            orientation: 'vertical'
                            BoxLayout:
                                size_hint: 0.9, 1.4
                                Label:
                                    text: '     X2 c накоплением'
                                    pos_hint: {"x":1, "y":0.3}
                                    text_size: self.size
                                    size_hint_x: 0.8
                                    font_size: 14
                                    color: 0.3,0.3,0.3,1
                                    valign: 'center'
                                ToggleButton:
                                    id: x2n
                                    size_hint_x: 0.23
                                    size_hint_y: 1.5
                                    group: 'b'
                                    state: 'down'
                                    background_normal: 'Emina/Enable/Option'
                                    background_down: 'Emina/Activated/Option'
                                    #state: 'normal' if x1s.active or x1n.active else 'down'
                            BoxLayout:
                                padding: [0,0,24,0]
                                Label:
                                    text: ' нужен excel-файл?'
                                    valign: 'top'
                                    halign: 'right'
                                    #size_hint_x: 0.7
                                    color: 0.3,0.3,0.3,1
                                ToggleButton:
                                    id: xl
                                    size_hint: 0.28, 1.4
                                    background_disabled_normal: 'Emina/Disable/Tumlr'
                                    background_disabled_down: 'Emina/Disable/Tumlr'
                                    background_normal: 'Emina/Enable/Tumlr'
                                    background_down: 'Emina/Activated/Tumlr'
                                    disabled: False if x2n.state == 'down' else True
                        BoxLayout:
                            orientation: 'vertical'
                            BoxLayout:
                                Label:
                                    text: '             X1 с накоплением'
                                    pos_hint: {"x":1, "y":0.3}
                                    text_size: self.size
                                    size_hint_x: 0.8
                                    font_size: 14
                                    color: 0.3,0.3,0.3,1
                                    valign: 'center'
                                ToggleButton:
                                    id: x1n
                                    size_hint_x: 0.205
                                    size_hint_y: 1.7
                                    group: 'b'
                                    background_normal: 'Emina/Enable/Option'
                                    background_down: 'Emina/Activated/Option'
                            BoxLayout:
                                Label:
                                    text: '     значения по бокам?'
                                    color: 0.3,0.3,0.3,1
                                ToggleButton:
                                    id: xo
                                    size_hint: 0.25, 1.17
                                    background_disabled_normal: 'Emina/Disable/Tumlr'
                                    background_disabled_down: 'Emina/Disable/Tumlr'
                                    background_normal: 'Emina/Enable/Tumlr'
                                    background_down: 'Emina/Activated/Tumlr'
                                    disabled: False if x1n.state == 'down' else True
                        BoxLayout:
                            orientation: 'vertical'
                            BoxLayout:
                                size_hint_x: 1.1
                                Label:
                                    text: '               X1 структурированная'
                                    pos_hint: {"x":1, "y":0.3}
                                    text_size: self.size
                                    size_hint_x: 0.9
                                    font_size: 14
                                    color: 0.3,0.3,0.3,1
                                    valign: 'center'
                                ToggleButton:
                                    id: x1s
                                    size_hint_x: 0.2
                                    size_hint_y: 1.7
                                    group: 'b'
                                    background_normal: 'Emina/Enable/Option'
                                    background_down: 'Emina/Activated/Option'
                            BoxLayout:
                    BoxLayout:
                        size_hint: None, None
                        size: (0, 30)
                    BoxLayout:
                        size_hint: None, None
                        size: (800, 100)
                        BoxLayout:
                            orientation: 'vertical'
                            size_hint: None, None
                            size: 395, 100
                            BoxLayout:
                                BoxLayout:
                                    size_hint_x: 0.04
                                Button:
                                    #disabled: False if inn_check.active else True
                                    text: "Сохранить график и xlsx" if xl.state == 'down' else "Сохранить график"
                                    background_normal: 'Emina/Enable/Button_b'
                                    background_down: 'Emina/Activated/Button_b'
                                    on_release: app.Action()
                            BoxLayout:
                                size_hint_y: 0.01
                        BoxLayout:
                            orientation: 'vertical'
                            size_hint_x: 0.43
                            canvas:
                                Rectangle:
                                    pos: (397, 6)
                                    size: (399, 100) # (399, 93)
                                    source: 'Emina/Enable/Default_5'
                            BoxLayout:
                                orientation: 'horizontal'
                                BoxLayout:
                                    orientation: 'vertical'
                                    size_hint: 2, 1
                                    BoxLayout:
                                        Label:
                                            text: '        Измерение'
                                            pos_hint: {"x":0, "y":0.03}
                                            text_size: self.size
                                            color: 0.3,0.3,0.3,1
                                            valign: 'center'
                                        DropBut:
                                            id: izm
                                            pos_hint: {"x":1, "y": -0.1}
                                            size_hint: None, None
                                            size: 150, 56
                                            text: 'в рублях'
                                            background_normal: 'Emina/Enable/Button_s'
                                            background_down: 'Emina/Activated/Button_s'
                                    BoxLayout:
                                        size_hint: 1.2, 0.9
                                        # padding: [0,0,0,20]
                                        Label:
                                            size_hint: 0.8, 0.9
                                            text: '                Делим?'
                                            pos_hint: {"x": 0, "y":0.24}
                                            text_size: self.size
                                            color: 0.3,0.3,0.3,1
                                            valign: 'center'
                                        ToggleButton:
                                            id: div
                                            size_hint: None, None
                                            size: 49, 31.6
                                            pos_hint: {"x": -0.2, "y":0.35}
                                            background_normal: 'Emina/Enable/Tumlr'
                                            background_down: 'Emina/Activated/Tumlr'
                                        BoxLayout:
                                BoxLayout:
                                    orientation: 'vertical'
                                    Label:
                                        text: '         Подписи?'
                                        size_hint: 1, 1.9
                                        text_size: self.size
                                        color: 0.3,0.3,0.3,1
                                        valign: 'center'
                                    ToggleButton:
                                        id: xp
                                        state: 'down'
                                        size_hint: None, None
                                        size: 49, 31.6
                                        pos_hint: {"x":0.34, "y":1}
                                        background_normal: 'Emina/Enable/Tumlr'
                                        background_down: 'Emina/Activated/Tumlr'
                                    BoxLayout:
            Screen:
                name: 'Info'
                ScrollView:
                    do_scroll_x: False
                    BoxLayout:
                        orientation: 'vertical'
                        size_hint_y: None
                        height: dp(1800)
                        Label:
                            text: 'Справка'
                            font_size: 48
                        BoxLayout:###################################
                            size_hint_y: 0.125
                            canvas:
                                Color:
                                    rgba: 1, 1, 1, 0.1
                                Rectangle:
                                    pos: self.pos
                                    size: self.size
                        Label:
                            text: 'Программа позволяет строить три диаграммы:'
                            font_size: 20
                            size_hint_y: 0.3
                        BoxLayout:
                            size_hint_y: 13
                            orientation: 'vertical'
                            BoxLayout:
                                orientation: 'horizontal'
                                Label:
                                    text: 'Двойная с накоплением'
                                    size_hint_x: 0.4
                                Image:
                                    source: "Emina/Plot1"
                                    mipmap: True
                            BoxLayout:
                                orientation: 'horizontal'
                                Label:
                                    text: 'Одинарная с накоплением'
                                    size_hint_x: 0.4
                                Image:
                                    source: "Emina/Plot2"
                                    mipmap: True
                            BoxLayout:
                                orientation: 'horizontal'
                                Label:
                                    text: 'Одинарная структурированная'
                                    size_hint_x: 0.4
                                Image:
                                    source: "Emina/Plot3"
                                    mipmap: True
                        BoxLayout:###################################
                            size_hint_y: 0.125
                            canvas:
                                Color:
                                    rgba: 1, 1, 1, 0.1
                                Rectangle:
                                    pos: self.pos
                                    size: self.size
                        BoxLayout:
                            size_hint_y: 5
                            BoxLayout:
                                orientation: 'vertical'
                                Label:
                                    text: 'Перетащите excel, где на первом листе будет'
                                    size_hint_y: 0.1
                                Label:
                                    text: 'сводная таблица со следующими настройками:'
                                    size_hint_y: 0.1
                                Image:
                                    source: "Emina/example"
                                    size_hint_x: 1
                            Image:
                                source: "Emina/setts"
                                size_hint_x: 0.6
                        BoxLayout:###################################
                            size_hint_y: 0.125
                            canvas:
                                Color:
                                    rgba: 1, 1, 1, 0.1
                                Rectangle:
                                    pos: self.pos
                                    size: self.size
                        Label:
                            text: "Желательно, чтобы в книге excel было до трех листов"
                            size_hint_y: 0.3
                        Label:
                            text: "(сводную таблицу можно сразу делать в новой книге)"
                            size_hint_y: 0.3
                        BoxLayout:###################################
                            size_hint_y: 0.125
                            canvas:
                                Color:
                                    rgba: 1, 1, 1, 0.1
                                Rectangle:
                                    pos: self.pos
                                    size: self.size
                        Label:
                            text: "Если требуется все же в ручную построить <Двойная с накоплением>,"
                            size_hint_y: 0.3
                        Label:
                            text: "то для этого можно поставить галочку в графе <Нужен excel-файл?>"
                            size_hint_y: 0.3
                        BoxLayout:###################################
                            size_hint_y: 0.125
                            canvas:
                                Color:
                                    rgba: 1, 1, 1, 0.1
                                Rectangle:
                                    pos: self.pos
                                    size: self.size
                        Label:
                            text: "Если требуется чтобы в <Одинарная с накоплением>"
                            size_hint_y: 0.3
                        Label:
                            text: "справедливая стоимость отражалась справа - отметьте <по бокам>"
                            size_hint_y: 0.3
                        BoxLayout:###################################
                            size_hint_y: 0.125
                            canvas:
                                Color:
                                    rgba: 1, 1, 1, 0.1
                                Rectangle:
                                    pos: self.pos
                                    size: self.size
                        Label:
                            text: "Также есть возможность указать размерность и при необходимости разделить"
                            size_hint_y: 0.4
                        BoxLayout:###################################
                            size_hint_y: 0.125
                            canvas:
                                Color:
                                    rgba: 1, 1, 1, 0.1
                                Rectangle:
                                    pos: self.pos
                                    size: self.size
                        Label:
                            text: "Все, что сверх - допиливается руками в MS Paint / PowerPoint"
                            size_hint_y: 0.4
                        BoxLayout:###################################
                            size_hint_y: 0.125
                            canvas:
                                Color:
                                    rgba: 1, 1, 1, 0.1
                                Rectangle:
                                    pos: self.pos
                                    size: self.size
                FloatLayout:
                    halign: 'right'
                    BoxLayout:
                        size_hint: None, None
                        size: dp(60), dp(60)
                        pos_hint:{'top': 1, 'right':1}
                        halign: 'right'
                        RoundedButton:
                            halign: 'right'
                            # size_hint:(0.1, 0.1)
                            # pos_hint:{'x': 0.925, 'y': 0.895}
                            background_color: 0,0,0,0
                            text: '<'
                            font_size: 50
                            canvas.before:
                                Color:
                                    rgba: (.3,.0,.9,.4) if self.state=='normal' else (1,1,1,1)
                                RoundedRectangle:
                                    pos: self.pos
                                    size: 60,60
                                    radius: [50,]
                            on_release:
                                scr_mngr.transition.direction = 'right'
                                scr_mngr.current = 'SP'
    TabbedPanelItem:
        text: 'Выпискатор'
        id: tab2
        #disabled: True
        BoxLayout:
            orientation: 'vertical'
            BoxLayout:
                orientation: 'vertical'
                size_hint_y:0.6
                BoxLayout:
                    orientation: 'horizontal'
                    Label:
                        text: "ВЫПИСКАТОР"
                        size_hint_x: 1
                        text_size: self.size
                        halign: 'left'
                        valign: 'bottom'
                        font_size: 54
                    BoxLayout:
                        size_hint: None, None
                        size: dp(60), dp(60)
                        pos_hint:{'top': 1, 'right':1}
                        RoundedButton:
                            background_color: 0,0,0,0
                            halign: 'right'
                            text: "?"
                            font_size: 50
                            on_release:
                                scr_mngr.transition.direction = 'left'
                                scr_mngr.current = 'Info'
                            canvas.before:
                                Color:
                                    rgba: (.3,.0,.9,1) if self.state=='normal' else (1,1,1,1)
                                RoundedRectangle:
                                    pos: self.pos
                                    size: 60,60
                                    radius: [50,]
                Label:
                    text: "Интерактивная выписка"
                    text_size: self.size
                    halign: 'right'
                    font_size: 32
            BoxLayout:
                size_hint_y: 0.03
                canvas:
                    Color:
                        rgba: 1, 1, 1, 0.1
                    Rectangle:
                        pos: self.pos
                        size: self.size
            BoxLayout:
                on_dropfile:
                Label:
                    text: "Drag'n Drop"
            BoxLayout:
                size_hint_y: 0.03
                canvas:
                    Color:
                        rgba: 1, 1, 1, 0.1
                    Rectangle:
                        pos: self.pos
                        size: self.size
            BoxLayout:
                orientation: 'vertical'
                size_hint_y: 0.2
                BoxLayout
                    # горизонтальный
                    Label:
                        text: "Загружено:"
                        halign: 'left'
                        size_hint_x: 0.4
                    BoxLayout:
                        # горизонтальный
                        size_hint_x: 0.2
                        Label:
                            text: "Массив"
                            #size_hint_x: 0.3
                        CheckBox:
                            id: massiv_ready
                            disabled: True
                            group: 'a'
                            # кружок с галочкой состояния загрузки файла
                            #size_hint_x: 0.1
                    BoxLayout:
                        # горизонтальный
                        size_hint_x: 0.7
                        Label:
                            text: "Заемщики"
                        CheckBox:
                            id: inn_ready
                            size_hint_x: 0.1
                            disabled: True
                            group: 'b'
                            # кружок с галочкой состояния загрузки файла
                        BoxLayout:
                        BoxLayout:
        #### ШАПКА ###
            BoxLayout:
                size_hint_y: 0.25
                BoxLayout:
                    orientation: 'vertical'
                    #size_hint_x: 0.2
                    Label:
                        text: "ИНН цели"
                        font_size: 9
                    TextInput:
                        id:TInn
                        text: 'A'
                BoxLayout:
                    orientation: 'vertical'
                    #size_hint_x: 0.2
                    Label:
                        text: "Цель"
                        font_size: 9
                    TextInput:
                        id:Target
                        text: 'B'
                BoxLayout:
                    orientation: 'vertical'
                    #size_hint_x: 0.2
                    Label:
                        text: "Счет цели"
                        font_size: 9
                    TextInput:
                        id:TAccount
                        text: 'C'
                BoxLayout:
                    orientation: 'vertical'
                    #size_hint_x: 0.2
                    Label:
                        text: "Дата"
                        font_size: 9
                    TextInput:
                        id:Data
                        text: 'D'
                BoxLayout:
                    orientation: 'vertical'
                    #size_hint_x: 0.2
                    Label:
                        text: "Банк"
                        font_size: 9
                    TextInput:
                        id:Corrbank
                        text: 'G'
                BoxLayout:
                    orientation: 'vertical'
                    #size_hint_x: 0.2
                    Label:
                        text: "Контрагент"
                        font_size: 9
                    TextInput:
                        id:Partner
                        text: 'J'
                BoxLayout:
                    orientation: 'vertical'
                    #size_hint_x: 0.2
                    Label:
                        text: "ИНН контрагента"
                        font_size: 9
                    TextInput:
                        id:Inn
                        text: 'K'
                BoxLayout:
                    orientation: 'vertical'
                    #size_hint_x: 0.2
                    Label:
                        text: "Счет контрагента"
                        font_size: 9
                    TextInput:
                        id:Account
                        text: 'L'
                BoxLayout:
                    orientation: 'vertical'
                    #size_hint_x: 0.2
                    Label:
                        text: "Дебет"
                        font_size: 9
                    TextInput:
                        id:Debet
                        text: 'N'
                BoxLayout:
                    orientation: 'vertical'
                    #size_hint_x: 0.2
                    Label:
                        text: "Кредит"
                        font_size: 9
                    TextInput:
                        id:Credit
                        text: 'O'
                BoxLayout:
                    orientation: 'vertical'
                    #size_hint_x: 0.2
                    Label:
                        text: "Назначение"
                        font_size: 9
                    TextInput:
                        id:Description
                        text: 'R'
        #### ШАПКА ###
            BoxLayout:
                size_hint_y: 0.15
                TextInput:
                    id: targetInn
                    hint_text: "Введите ИНН"
                    #text: app.INNTARGET
                DropBut:
                    # Список ИНН, подтягиваемый с заемщиков DropDown
                    size_hint: 0.37, 1
                    text: 'выбор контрагента'
                    # canvas:
                    #     Rectangle:
                    #         #source: self.icon
                    #         pos: self.center_x-25, self.center_y-25
                    #         size: 50, 50
                CheckBox:
                    id: inn_check
                    size_hint_x: 0.1
                    active: True if massiv_ready.active and targetInn.text != '' else False
                    disabled: True
                    # галочка со статусом
            BoxLayout:
                size_hint_y: 0.1
            Button:
                size_hint_y: 0.3
                disabled: False if inn_check.active else True
                text: "Визуализация"
                on_release: app.Visualize()
    TabbedPanelItem:
        id: tab3
        text:'Сбиватор'
        FloatLayout:
            Button:
                size_hint_x: 0.5
                on_release: app.change_Vid()
                Video:
                    size: (800, 574)
                    id: video
                    vid: '0'
                    source: './Emina/' + self.vid
                    state: 'play'
            BoxLayout:
                size_hint: None, None
                size: dp(60), dp(60)
                pos_hint:{'top': 1, 'right':1}
                RoundedButton:
                    background_color: 0,0,0,0
                    halign: 'right'
                    text: "<"
                    font_size: 50
                    on_release:
                        app.change_Vid(reset=1)
                    canvas.before:
                        Color:
                            rgba: (.3,.0,.9,.4) if self.state=='normal' else (1,1,1,1)
                        RoundedRectangle:
                            pos: self.pos
                            size: 60,60
                            radius: [50,]
            BoxLayout:
                id: q1
                orientation: 'vertical'
                size_hint: (None, None)
                height: dp(150)
                width: dp(100)
                pos_hint: {'x':0.584, 'y': 0.39}
                opacity: 0
                BoxLayout:
                    orientation: 'vertical'
                    size_hint_y: None
                    height: dp(90)
                    Label:
                        font_size: 50
                        color: (69/255,255/255,66/255,1)
                        text: 'ЧЕГО'
                    Label:
                        font_size: 30
                        color: (69/255,255/255,66/255,1)
                        text: 'ТЫ'
                    Label:
                        font_size: 30
                        color: (69/255,255/255,66/255,1)
                        text: '  ХОЧЕШЬ'
                BoxLayout:
                    size_hint_y: 0.1
                BoxLayout:
                    orientation: 'horizontal'
                    size_hint_y: 0.8
                    Button:
                        font_size: 20
                        size_hint_y: 0.8
                        text: 'ВПР'
                        on_release: app.change_Vid()
                    BoxLayout:
                        orientation: 'vertical'
                        Label:
                            text: 'точность'
                            size_hint_x: 1.3
                            halign: 'right'
                        BoxLayout:
                            orientation: 'horizontal'
                            TextInput:
                                id: vpr_acc
                                text: '100'
                            Label:
                                size_hint_x: 0.3
                                text: '%'
            BoxLayout:
                id: hint1
                orientation: 'vertical'
                size_hint: (0.2, None)
                pos_hint: {'x':0.56, 'y': 0.2}
                height: dp(90)
                opacity: 0
                Label:
                    font_size: 45
                    color: (69/255,255/255,66/255,1)
                    text: 'Ctrl+C'
                Label:
                    font_size: 30
                    color: (69/255,255/255,66/255,1)
                    text: 'к чему'
                Label:
                    font_size: 30
                    color: (69/255,255/255,66/255,1)
                    text: 'тянем'
            BoxLayout:
                id: hint2
                orientation: 'vertical'
                size_hint:(0.3, None)
                height: dp(90)
                pos_hint: {'x':0.59, 'y': 0.45}
                opacity: 0
                Label:
                    font_size: 50
                    color: (69/255,255/255,66/255,1)
                    text: 'Ctrl+C'
                Label:
                    font_size: 30
                    color: (69/255,255/255,66/255,1)
                    text: 'что тянем'
                BoxLayout:
                    BoxLayout:
                        size_hint_x: 0.3
                    BoxLayout:
                        BoxLayout:
                            BoxLayout:
                                orientation: 'vertical'
                                size_hint_x: 0.3
                                Label:
                                    font_size: 10
                                    text: 'по какой'
                                Label:
                                    font_size: 10
                                    text: 'колонке'
                            TextInput:
                                id: vpr_col1
                                size_hint_x: 0.3
                                text: '1'
                        BoxLayout:
                            opacity: 0 if vpr_acc.text == '100' else 1
                            disabled: True if vpr_acc.text == '100' else False
                            BoxLayout:
                                orientation: 'vertical'
                                size_hint_x: 0.3
                                Label:
                                    font_size: 10
                                    text: 'какую'
                                Label:
                                    font_size: 10
                                    text: 'колонку'
                            TextInput:
                                id: vpr_col2
                                size_hint_x: 0.3
                                text: '2'
                        BoxLayout:
                            size_hint_x: 0.2
            BoxLayout:
                id: hint3
                orientation: 'vertical'
                size_hint: (0.2, None)
                pos_hint: {'x':0.66, 'y': 0.4}
                height: dp(90)
                opacity: 0
                Label:
                    font_size: 35
                    color: (69/255,255/255,66/255,1)
                    text: 'СМОТРИ'
                Label:
                    font_size: 30
                    color: (69/255,255/255,66/255,1)
                    text: 'РАБОЧИЙ'
                Label:
                    font_size: 30
                    color: (69/255,255/255,66/255,1)
                    text: 'СТОЛ'
    TabbedPanelItem:
        id: tab4
        text:'Редактор'
        ScreenManager:
            id: scr_mngr2
            Screen:
                name: 'Menu'
                BoxLayout:
                    id: main
                    orientation: 'vertical'
                    canvas.before:
                        Rectangle:
                            size: self.size
                            source: 'Emina/Enable/Default'
                    BoxLayout:
                        orientation: 'vertical'
                        size_hint_y: 0.4
                        canvas.before:
                            Rectangle:
                                pos: (15, 489)
                                size: (324, 81)
                                source: 'Emina/Enable/Default_1'
                        BoxLayout:
                            Label:
                                text: "   РЕДАКТОР"
                                size_hint_x: 11.8
                                pos_hint: {"x":1, "y":-0.2}
                                text_size: self.size
                                halign: 'left'
                                valign: 'bottom'
                                font_size: 48
                                color: 0.3,0.3,0.3,1
                            BoxLayout:
                                size_hint: None, None
                                size: dp(85), dp(85)
                                pos_hint:{'top': 1.06, 'right': 1}
                                RoundedButton:
                                    background_normal: 'Emina/Enable/Help'
                                    background_down: 'Emina/Activated/Help'
                                    halign: 'right'
                                    text: "?"
                                    font_size: 50
                                    pos: (1,0.5)
                                    on_release:
                                        scr_mngr2.transition.direction = 'left'
                                        scr_mngr2.current = 'Info2'
                        Label:
                            text: "Типограф для отчётов  "
                            text_size: self.size
                            halign: 'right'
                            color: 0.3,0.3,0.3,1
                            font_size: 32
                            canvas.before:
                                Rectangle:
                                    pos: (427, 448)
                                    size: (367, 56)
                                    source: 'Emina/Enable/Default_2'
                        BoxLayout:
                            size_hint_y: 0.3
                    BoxLayout:
                        on_dropfile:
                        orientation: 'vertical'
                        canvas.before:
                            Rectangle:
                                pos: (290, 215)
                                size: (217, 223)
                                source: 'Emina/Enable/Default_3'
                        BoxLayout:
                        Label:
                            size_hint_y: 0.2
                            text: "Drag'n Drop"
                            color: 0.3,0.3,0.3,1
                        Image:
                            id: doc_img
                            source: 'Emina/Enable/DnD'
                        BoxLayout:
                            size_hint_y: 2.7
                    BoxLayout:
                        size_hint: None, None
                        size: (800, 100)
                        BoxLayout:
                            orientation: 'vertical'
                            size_hint: None, None
                            size: 395, 100
                            BoxLayout:
                                BoxLayout:
                                    size_hint_x: 0.04
                                Button:
                                    #disabled: False if inn_check.active else True
                                    text: "Исправить отчёт"
                                    background_normal: 'Emina/Enable/Button_b'
                                    background_down: 'Emina/Activated/Button_b'
                                    on_release: app.Typograph()
                            BoxLayout:
                                size_hint_y: 0.01
                        BoxLayout:
                            orientation: 'vertical'
                            size_hint_x: 0.2
                            canvas.before:
                                Rectangle:
                                    pos: (390, 6)
                                    size: (200, 98)
                                    source: 'Emina/Enable/Default_6'
                            BoxLayout:
                                orientation: 'vertical'
                                Label:
                                    text: '        Исправлять даты?'
                                    text_size: self.size
                                    color: 0.3,0.3,0.3,1
                                    valign: 'center'
                                ToggleButton:
                                    id: ndate
                                    size_hint: None, None
                                    size: 49, 31.6
                                    state: 'down'
                                    pos_hint: {"x":0.34, "y":1}
                                    background_normal: 'Emina/Enable/Tumlr'
                                    background_down: 'Emina/Activated/Tumlr'
                            BoxLayout:
                                size_hint_y: 0.3
                        BoxLayout:
                            orientation: 'vertical'
                            size_hint_x: 0.2
                            canvas.before:
                                Rectangle:
                                    pos: (590, 6)
                                    size: (200, 98)
                                    source: 'Emina/Enable/Default_6'
                            BoxLayout:
                                orientation: 'vertical'
                                Label:
                                    text: '               Руб. на ₽?'
                                    text_size: self.size
                                    color: 0.3,0.3,0.3,1
                                    valign: 'center'
                                ToggleButton:
                                    id: rubl
                                    state: 'normal'
                                    size_hint: None, None
                                    size: 49, 31.6
                                    pos_hint: {"x":0.34, "y":1}
                                    background_normal: 'Emina/Enable/Tumlr'
                                    background_down: 'Emina/Activated/Tumlr'
                            BoxLayout:
                                size_hint_y: 0.3
            Screen:
                name: 'Info2'
                BoxLayout:
                    orientation: 'vertical'
                    #size_hint_y: None
                    Label:
                        text: 'Справка'
                        font_size: 48
                    BoxLayout:###################################
                        size_hint_y: 0.125
                        canvas:
                            Color:
                                rgba: 1, 1, 1, 0.1
                            Rectangle:
                                pos: self.pos
                                size: self.size
                    Label:
                        text: 'Требуется перетащить файл Word в окно программы.'
                        font_size: 20
                        size_hint_y: 0.3
                    BoxLayout:###################################
                        size_hint_y: 0.125
                        canvas:
                            Color:
                                rgba: 1, 1, 1, 0.1
                            Rectangle:
                                pos: self.pos
                                size: self.size
                    Label:
                    Label:
                        text: 'Программа поправит очевидные опечатки, расставит'
                        font_size: 16
                        size_hint_y: 0.3
                    Label:
                        text: 'нормальные кавычки («елочки»), уберет лишние пробелы,'
                        font_size: 16
                        size_hint_y: 0.3
                    Label:
                        text: 'в нужных местах поменяет дефисы на тире;'
                        font_size: 16
                        size_hint_y: 0.3
                    Label:
                        text: 'неразрывным пробелом соединит: числа с размерностью,'
                        font_size: 16
                        size_hint_y: 0.3
                    Label:
                        text: 'организационно-правовую форму с названием компании и'
                        font_size: 16
                        size_hint_y: 0.3
                    Label:
                        text: 'поправит формат дат.'
                        font_size: 16
                        size_hint_y: 0.3
                    BoxLayout:
                        #size_hint_y: dp(960)
                FloatLayout:
                    halign: 'right'
                    BoxLayout:
                        size_hint: None, None
                        size: dp(60), dp(60)
                        pos_hint:{'top': 1, 'right':1}
                        halign: 'right'
                        RoundedButton:
                            halign: 'right'
                            # size_hint:(0.1, 0.1)
                            # pos_hint:{'x': 0.925, 'y': 0.895}
                            background_color: 0,0,0,0
                            text: '<'
                            font_size: 50
                            canvas.before:
                                Color:
                                    rgba: (.3,.0,.9,.4) if self.state=='normal' else (1,1,1,1)
                                RoundedRectangle:
                                    pos: self.pos
                                    size: 60,60
                                    radius: [50,]
                            on_release:
                                scr_mngr2.transition.direction = 'right'
                                scr_mngr2.current = 'Menu'
    TabbedPanelItem:
        id: tab5
        text:'Недвигер'
        ScreenManager:
            id: scr_mngr3
            Screen:
                name: 'Menu'
                BoxLayout:
                    id: main
                    orientation: 'vertical'
                    canvas.before:
                        Rectangle:
                            size: self.size
                            source: 'Emina/Enable/Default'
                    BoxLayout:
                        orientation: 'vertical'
                        size_hint_y: 0.4
                        canvas.before:
                            Rectangle:
                                pos: (15, 489)
                                size: (324, 81)
                                source: 'Emina/Enable/Default_1'
                        BoxLayout:
                            Label:
                                text: "   НЕДВИГЕР"
                                size_hint_x: 11.8
                                pos_hint: {"x":1, "y":-0.2}
                                text_size: self.size
                                halign: 'left'
                                valign: 'bottom'
                                font_size: 48
                                color: 0.3,0.3,0.3,1
                            BoxLayout:
                                size_hint: None, None
                                size: dp(85), dp(85)
                                pos_hint:{'top': 1.06, 'right': 1}
                                RoundedButton:
                                    background_normal: 'Emina/Enable/Help'
                                    background_down: 'Emina/Activated/Help'
                                    halign: 'right'
                                    text: "?"
                                    font_size: 50
                                    pos: (1,0.5)
                                    on_release:
                                        scr_mngr3.transition.direction = 'left'
                                        scr_mngr3.current = 'Info3'
                        Label:
                            text: "БД по переходам прав "
                            text_size: self.size
                            halign: 'right'
                            color: 0.3,0.3,0.3,1
                            font_size: 32
                            canvas.before:
                                Rectangle:
                                    pos: (427, 448)
                                    size: (367, 56)
                                    source: 'Emina/Enable/Default_2'
                        BoxLayout:
                            size_hint_y: 0.3
                    BoxLayout:
                        on_dropfile:
                        orientation: 'vertical'
                        canvas.before:
                            Rectangle:
                                pos: (290, 215)
                                size: (217, 223)
                                source: 'Emina/Enable/Default_3'
                        BoxLayout:
                        Label:
                            size_hint_y: 0.2
                            text: "Drag'n Drop"
                            color: 0.3,0.3,0.3,1
                        Image:
                            id: doc_img
                            source: 'Emina/Enable/DnD'
                        BoxLayout:
                            size_hint_y: 2.7
                    BoxLayout:
                        size_hint: None, None
                        size: (800, 100)
                        BoxLayout:
                            orientation: 'vertical'
                            size_hint: None, None
                            size: 395, 100
                            BoxLayout:
                                BoxLayout:
                                    size_hint_x: 0.04
                                Button:
                                    #disabled: False if inn_check.active else True
                                    text: "Сохранить базу данных"
                                    background_normal: 'Emina/Enable/Button_b'
                                    background_down: 'Emina/Activated/Button_b'
                                    on_release: app.EGRN()
                            BoxLayout:
                                size_hint_y: 0.01
            Screen:
                name: 'Info3'
                BoxLayout:
                    orientation: 'vertical'
                    #size_hint_y: None
                    Label:
                        text: 'Справка'
                        font_size: 48
                    BoxLayout:###################################
                        size_hint_y: 0.125
                        canvas:
                            Color:
                                rgba: 1, 1, 1, 0.1
                            Rectangle:
                                pos: self.pos
                                size: self.size
                    Label:
                        text: 'Требуется перетащить папку с .xml файлами в окно программы'
                        font_size: 20
                        size_hint_y: 0.3
                    BoxLayout:###################################
                        size_hint_y: 0.125
                        canvas:
                            Color:
                                rgba: 1, 1, 1, 0.1
                            Rectangle:
                                pos: self.pos
                                size: self.size
                    Label:
                    Label:
                        text: 'Программа пройдется по каждому файлу,'
                        font_size: 16
                        size_hint_y: 0.3
                    Label:
                        text: 'вытащит все содержащиеся данные и '
                        font_size: 16
                        size_hint_y: 0.3
                    Label:
                        text: 'сохранит excel-файл на рабочем столе.'
                        font_size: 16
                        size_hint_y: 0.3
                    BoxLayout:
                        #size_hint_y: dp(960)
                FloatLayout:
                    halign: 'right'
                    BoxLayout:
                        size_hint: None, None
                        size: dp(60), dp(60)
                        pos_hint:{'top': 1, 'right':1}
                        halign: 'right'
                        RoundedButton:
                            halign: 'right'
                            # size_hint:(0.1, 0.1)
                            # pos_hint:{'x': 0.925, 'y': 0.895}
                            background_color: 0,0,0,0
                            text: '<'
                            font_size: 50
                            canvas.before:
                                Color:
                                    rgba: (.3,.0,.9,.4) if self.state=='normal' else (1,1,1,1)
                                RoundedRectangle:
                                    pos: self.pos
                                    size: 60,60
                                    radius: [50,]
                            on_release:
                                scr_mngr3.transition.direction = 'right'
                                scr_mngr3.current = 'Menu'

<Popup>:
    id: err_popup
    title:     "Ошибка"
    size_hint_y: 0.9
    size_hint_x: 0.9
    separator_color: 0,0,0,0
    title_color: 0.3,0.3,0.3,1
    title_size: 18
    background: 'Emina/Enable/Default'
    BoxLayout:
        orientation: "vertical"
        Label:
            text: app.popup_errtxt()[0]
            size_hint_y: 0.2
            color: 0.3,0.3,0.3,1
            canvas.before:
                Rectangle:
                    pos: (15, 489)
                    size: (324, 81)
                    source: 'Emina/Enable/Default_1'
        Image:
            id: image
            source: "Emina/setts"
            opacity: app.popup_errtxt()[1]
        Button:
            size_hint:  (0.4, 0.2)
            text: "Вернуться"
            background_normal: 'Emina/Enable/Button_s'
            background_down: 'Emina/Activated/Button_s'
            on_release: err_popup.dismiss()

<RoundedButton@Button>:
    # Кнопка справки
    # text: '?'
    # font_size: 50
    # on_release:
    #     app.root.transition.direction = 'left'
    #     app.root.current = 'Info'
'''


inn_table = ['Всех Заемщиков', '', '']
inns = None
file_paths = {'Data': None, 'Inn': None}
types = ListProperty(['Всех Заемщиков'])
event = None


class SmartPlot(App):
    file_paths = {'Data': None, 'Inn': None}
    inn_table = ['Всех Заемщиков', '1', '2']
    types = ListProperty(['Всех Заемщиков'])
    inns = None
    vid_state = 1
    vid_list = ['0', '1_intro', '2_hand_L', '2-1_error', '3_hand_R', '3-1_error', '4_calc', '4-1_error', '5_after']
    df1 = None
    df2 = None
    df3 = None
    acc = None
    event = None

    def __init__(self, **kwargs):
        super(SmartPlot, self).__init__(**kwargs)

    def build(self):
        main_widget = Builder.load_string(main_widget_kv)
        Window.bind(on_dropfile=self._on_file_drop)
        self.icon = 'icon.ico'
        return main_widget

    def _on_file_drop(self, window, file_path):
        inn_table = 123
        path = file_path.decode('utf-8')
        if App.get_running_app().root.current_tab.text == 'Выпискатор':  # если выбран таб Выпискатора:
            try:
                if re.search('ИНН', path).group() == 'ИНН':
                    App.get_running_app().root.ids.inn_ready.active = True
                    SmartPlot.file_paths['Inn'] = path
                    try:
                        SmartPlot.inns = pd.read_excel(SmartPlot.file_paths['Inn'], header=None, skiprows=1)
                        print(0)
                        if len(SmartPlot.inns.columns) > 1:
                            print(1)
                            SmartPlot.inn_table = sorted(SmartPlot.inns[1].astype(str).tolist())
                        else:
                            print(2)
                            SmartPlot.inn_table = sorted(SmartPlot.inns[0].astype(str).tolist())
                    except Exception as e:
                        print(e)
                        SmartPlot.inn_table = None
            except Exception as e:
                print(e)
                App.get_running_app().root.ids.massiv_ready.active = True
                SmartPlot.file_paths['Data'] = path
        # file_path = None
        elif App.get_running_app().root.current_tab.text == 'Сбиватор':  # если выбран таб Сбиватора:
            if SmartPlot.vid_state == 3:
                SmartPlot.df1 = pd.read_excel(path, header=None)
            elif SmartPlot.vid_state == 5:
                SmartPlot.df2 = pd.read_excel(path, header=None)
        else:
            SmartPlot.file_paths = file_path
            App.get_running_app().root.ids.doc_img.source = 'Emina/Activated/DnD'

    def popup_errtxt(self):
        if App.get_running_app().root.current_tab.text == "Редактор":
            txt = "Перетащите Word-файл с текстом в окно программы!"
            img_hide = 0
        elif App.get_running_app().root.current_tab.text == "Недвигер":
            txt = "Перетащите папку с .xml-файлами в окно программы!"
            img_hide = 0
        else:
            txt = "Перетащите Excel со сводной таблицей (настройки см. ниже) в окно программы!"
            img_hide = 1
        return [txt, img_hide]

    # ОСНОВНАЯ ФУНКЦИЯ SMARTPLOT
    def Action(self):
        if App.get_running_app().root.ids.x1s.state == 'down':
            n = 3
        elif App.get_running_app().root.ids.x1n.state == 'down':
            n = 2
        elif App.get_running_app().root.ids.x2n.state == 'down':
            n = 1
        else:
            n = 1
        delim = 1
        if App.get_running_app().root.ids.izm.text == 'в рублях':
            izmerenie = 'руб.'
        if App.get_running_app().root.ids.izm.text == 'в тысячах':
            izmerenie = 'тыс. руб.'
            if App.get_running_app().root.ids.div.state == 'down':
                delim = 1000
        if App.get_running_app().root.ids.izm.text == 'в миллионах':
            izmerenie = 'млн руб.'
            if App.get_running_app().root.ids.div.state == 'down':
                delim = 1000000
        if App.get_running_app().root.ids.izm.text == 'в миллиардах':
            izmerenie = 'млрд руб.'
            if App.get_running_app().root.ids.div.state == 'down':
                delim = 1000000000
        try:
            file = r'' + SmartPlot.file_paths.decode("utf-8")
            from pandas import read_excel, ExcelWriter
            from numpy import arange, array, nan_to_num, zeros, squeeze, divide, zeros_like
            df = read_excel(file, header=(1, 2), index_col=0)
            try:
                df = df / delim
                df = df.round()
            except Exception as e:
                print(e)
            # получаем 1 уровень колонок
            Netto = df.columns.get_level_values(0).unique()[0]
            Korka = df.columns.get_level_values(0).unique()[1]
            SS = df.columns.get_level_values(0).unique()[2]
            NettoSum = df.columns.get_level_values(0).unique()[3]
            NettoSumM = df[NettoSum].columns[0]
            SSSum = df.columns.get_level_values(0).unique()[5]

            # сортируем
            df = df.sort_values(by=[(NettoSum, NettoSumM)], ascending=False)

            # сохраняем ИТОГ
            SumNetto = df.iloc[0:1][NettoSum].values[0][0]
            SumSS = df.iloc[0:1][SSSum].values[0][0]

            # Удаляем общий итог
            df = df[1:]

            # делаем числовой индекс
            df = df.reset_index()

            # Здесь нужно оставить место для статьи баланса и для значений
            maxnetto = int(df[NettoSum][0:1].iloc[0])
            df['NettoS'] = maxnetto * 2 * 0.1
            df['SSS'] = maxnetto * 2 * 0.1
            df['IndeX'] = maxnetto * 2 * 0.3
            df['Dop'] = maxnetto - df[NettoSum]
            df['DopEnd'] = maxnetto - df[SSSum]

            DF = df[['Dop', Netto, 'NettoS', 'IndeX', 'SSS', SS, 'DopEnd']]

            # import os
            from os.path import expanduser as osexp
            if App.get_running_app().root.ids.xl.state == 'down':
                # writer = pd.ExcelWriter(os.path.expanduser('~/Desktop/ASVplot.xlsx'))
                writer = ExcelWriter(osexp('~/Desktop/ASVplot.xlsx'))
                DF.to_excel(writer, 'Группы')
                workbook = writer.book
                worksheet = writer.sheets['Группы']
                format1 = workbook.add_format({'num_format': '# ##0'})
                worksheet.set_column('B:Z', 12, format1)
                writer.save()

            # Пересортируем для графика
            df = df.sort_values(by=[(NettoSum, NettoSumM)], ascending=True)
            df = df.reset_index()
            df.loc[len(df)] = 0
            df['head1'] = 0
            df.loc[len(df) - 1, 'head1'] = maxnetto + maxnetto * 2 * 0.1
            df['head2'] = 0
            df.loc[len(df) - 1, 'head2'] = maxnetto * 2 * 0.3
            df['head3'] = 0
            df.loc[len(df) - 1, 'head3'] = maxnetto + maxnetto * 2 * 0.1
            DF = df[['head1', 'head2', 'head3', 'Dop', Netto, 'NettoS', 'IndeX', 'SSS', SS, 'DopEnd']]

            # перенос текста статей
            def Wrap(a):
                try:
                    if len(a) > 25:
                        txt = a
                        sps = []
                        j = 0
                        for i in range(10):
                            if j >= 0:
                                j = txt.find(' ', j + 1)
                                if j >= 0:
                                    sps.append(j)
                        for i in range(len(sps)):
                            if sps[-i - 1] <= 25:
                                iii = sps.index(sps[-i - 1])
                                ntxt = '%s%s%s' % (txt[0:sps[iii] + 1], '\n', txt[sps[iii] + 1:])
                                break
                        return (ntxt)
                    else:
                        return (a)
                except:
                    return (a)

            df['index'] = df['index'].map(lambda x: Wrap(x))
            # Понять сколько групп
            group_count = len(df[Netto].columns)
            from matplotlib.pyplot import rcParams, axis, barh, text as plttext, plot, legend, savefig, \
                close as pltclose
            from matplotlib.ticker import MultipleLocator
            from math import trunc  # для форматирования чисел

            # настройки визуала для Matplotlib
            ### всего статей
            y = arange(len(DF.index))
            ## динамический размер итоговой картинки в зависимости от числа статей
            plot_height = 1.7 + len(y[:-1]) * 0.83333333333333
            rcParams["figure.figsize"] = (20, plot_height)
            rcParams['figure.dpi'] = 500
            rcParams['savefig.dpi'] = 500
            rcParams["font.family"] = "Arial"
            axis('off')
            plot_font = {'family': 'Arial', 'size': '18', 'stretch': 'condensed'}
            plot_font_bold = {'family': 'Arial', 'size': '18', 'stretch': 'condensed', 'weight': 'bold'}
            plot_font_labl = {'family': 'Arial', 'size': '10', 'stretch': 'condensed'}

            # значения диаграмм по группам активов --------------------------------------------------------------------
            ## Подписи для легенды - заглушки
            greenlabel = None
            bluelabel = None
            graylabel = None
            yelowlabel = None
            orangelabel = None
            ## Определяем номер столбца для каждой из возможных групп активов
            spisok = df[Netto].columns
            for i in range(len(spisok)):
                if spisok[i].find('прод') >= 0:
                    green_netto_col = 4 + i
                    green_ss_col = 4 + i + len(spisok) - (i + 1) + 4 + i
                    greenlabel = 'К реализации'
                if spisok[i].find('Прод') >= 0:
                    green_netto_col = 4 + i
                    green_ss_col = 4 + i + len(spisok) - (i + 1) + 4 + i
                    greenlabel = 'К реализации'
                if spisok[i].find('реал') >= 0:
                    green_netto_col = 4 + i
                    green_ss_col = 4 + i + len(spisok) - (i + 1) + 4 + i
                    greenlabel = 'К реализации'
                if spisok[i].find('Реал') >= 0:
                    green_netto_col = 4 + i
                    green_ss_col = 4 + i + len(spisok) - (i + 1) + 4 + i
                    greenlabel = 'К реализации'
                if spisok[i].find('погаш') >= 0:
                    blue_netto_col = 4 + i
                    blue_ss_col = 4 + i + len(spisok) - (i + 1) + 4 + i
                    bluelabel = 'До погашения'
                if spisok[i].find('Погаш') >= 0:
                    blue_netto_col = 4 + i
                    blue_ss_col = 4 + i + len(spisok) - (i + 1) + 4 + i
                    bluelabel = 'До погашения'
                if spisok[i].find('ТДИ') >= 0:
                    gray_netto_col = 4 + i
                    gray_ss_col = 4 + i + len(spisok) - (i + 1) + 4 + i
                    graylabel = 'ТДИ'
                if spisok[i].find('тди') >= 0:
                    gray_netto_col = 4 + i
                    gray_ss_col = 4 + i + len(spisok) - (i + 1) + 4 + i
                    graylabel = 'ТДИ'
                if spisok[i].find('техн') >= 0:
                    yellow_netto_col = 4 + i
                    yellow_ss_col = 4 + i + len(spisok) - (i + 1) + 4 + i
                    yelowlabel = 'Техника'
                if spisok[i].find('Техн') >= 0:
                    yellow_netto_col = 4 + i
                    yellow_ss_col = 4 + i + len(spisok) - (i + 1) + 4 + i
                    yelowlabel = 'Техника'
                if spisok[i].find('спис') >= 0:
                    orange_netto_col = 4 + i
                    orange_ss_col = 4 + i + len(spisok) - (i + 1) + 4 + i
                    orangelabel = 'Списание'
                if spisok[i].find('Спис') >= 0:
                    orange_netto_col = 4 + i
                    orange_ss_col = 4 + i + len(spisok) - (i + 1) + 4 + i
                    orangelabel = 'Списание'

            if n == 1:
                ## Проставляем значения столбцов
                h1 = array(DF.iloc[:, 0])
                h2 = array(DF.iloc[:, 1])
                h3 = array(DF.iloc[:, 2])
                dop_one = nan_to_num(array(DF.Dop))
                try:
                    green_netto = nan_to_num(DF.iloc[:, green_netto_col]).astype('float')
                except:
                    green_netto = zeros(len(DF.iloc[:, 4]))
                try:
                    blue_netto = nan_to_num(array(DF.iloc[:, blue_netto_col])).astype('float')
                except:
                    blue_netto = zeros(len(DF.iloc[:, 4]))
                try:
                    gray_netto = nan_to_num(array(DF.iloc[:, gray_netto_col])).astype('float')
                except:
                    gray_netto = zeros(len(DF.iloc[:, 4]))
                try:
                    yellow_netto = nan_to_num(array(DF.iloc[:, yellow_netto_col])).astype('float')
                except:
                    yellow_netto = zeros(len(DF.iloc[:, 4]))
                try:
                    orange_netto = nan_to_num(array(DF.iloc[:, orange_netto_col])).astype('float')
                except:
                    orange_netto = zeros(len(DF.iloc[:, 4]))
                value_netto = nan_to_num(array(DF.iloc[:, len(spisok) + 4]))
                names = nan_to_num(array(DF.iloc[:, len(spisok) + 5]))
                value_ss = nan_to_num(array(DF.iloc[:, len(spisok) + 6]))
                try:
                    green_ss = nan_to_num(array(DF.iloc[:, green_ss_col])).astype('float')
                except:
                    green_ss = zeros(len(DF.iloc[:, 4]))
                try:
                    blue_ss = nan_to_num(array(DF.iloc[:, blue_ss_col])).astype('float')
                except:
                    blue_ss = zeros(len(DF.iloc[:, 4]))
                try:
                    gray_ss = nan_to_num(array(DF.iloc[:, gray_ss_col])).astype('float')
                except:
                    gray_ss = zeros(len(DF.iloc[:, 4]))
                try:
                    yellow_ss = nan_to_num(array(DF.iloc[:, yellow_ss_col])).astype('float')
                except:
                    yellow_ss = zeros(len(DF.iloc[:, 4]))
                try:
                    orange_ss = nan_to_num(array(DF.iloc[:, orange_ss_col])).astype('float')
                except:
                    orange_ss = zeros(len(DF.iloc[:, 4]))
                dop_two = nan_to_num(array(DF.DopEnd))

                # отрисовка диаграмм
                barh(y, h1, color=('#02003d'), height=1.2, alpha=0)
                barh(y, h2, color=('#02003d'), height=1.2, left=h1, alpha=0)
                barh(y, h3, color=('#02003d'), height=1.2, left=list(map(lambda h, hh: h + hh, h1, h2)), alpha=0)
                barh(y, dop_one, color='white', alpha=0)
                barh(y, green_netto, color=('#8cba88'), left=dop_one, label=greenlabel)
                barh(y, blue_netto, color=('#79a5d4'), label=bluelabel,
                     left=list(map(lambda do, gn: do + gn, dop_one, green_netto)))
                barh(y, gray_netto, color=('#878787'), label=graylabel,
                     left=list(map(lambda do, gn, bn: do + gn + bn, dop_one, green_netto, blue_netto)))
                barh(y, yellow_netto, color=('#f4e542'), label=yelowlabel, left=list(
                    map(lambda do, gn, bn, grn: do + gn + bn + grn, dop_one, green_netto, blue_netto, gray_netto)))
                barh(y, orange_netto, color=('#ff9191'), label=orangelabel, left=list(
                    map(lambda do, gn, bn, grn, yn: do + gn + bn + grn + yn, dop_one, green_netto, blue_netto,
                        gray_netto, yellow_netto)))
                barh(y, value_netto, color='white', alpha=0, left=list(
                    map(lambda do, gn, bn, grn, yn, on: do + gn + bn + grn + yn + on, dop_one, green_netto, blue_netto,
                        gray_netto, yellow_netto, orange_netto)))
                barh(y, names, color='white', alpha=0, left=list(
                    map(lambda do, gn, bn, grn, yn, on, vn: do + gn + bn + grn + yn + on + vn, dop_one, green_netto,
                        blue_netto, gray_netto, yellow_netto, orange_netto, value_netto)))
                barh(y, value_ss, color='white', alpha=0, left=list(
                    map(lambda do, gn, bn, grn, yn, on, vn, n: do + gn + bn + grn + yn + on + vn + n, dop_one,
                        green_netto, blue_netto, gray_netto, yellow_netto, orange_netto, value_netto, names)))
                barh(y, green_ss, color=('#8cba88'), left=list(
                    map(lambda do, gn, bn, grn, yn, on, vn, n, vs: do + gn + bn + grn + yn + on + vn + n + vs, dop_one,
                        green_netto, blue_netto, gray_netto, yellow_netto, orange_netto, value_netto, names, value_ss)))
                barh(y, blue_ss, color=('#79a5d4'), left=list(
                    map(lambda do, gn, bn, grn, yn, on, vn, n, vs, gs: do + gn + bn + grn + yn + on + vn + n + vs + gs,
                        dop_one, green_netto, blue_netto, gray_netto, yellow_netto, orange_netto, value_netto, names,
                        value_ss, green_ss)))
                barh(y, gray_ss, color=('#878787'), left=list(map(lambda do, gn, bn, grn, yn, on, vn, n, vs, gs,
                                                                         bs: do + gn + bn + grn + yn + on + vn + n + vs + gs + bs,
                                                                  dop_one, green_netto, blue_netto, gray_netto,
                                                                  yellow_netto, orange_netto, value_netto, names,
                                                                  value_ss, green_ss, blue_ss)))
                barh(y, yellow_ss, color=('#f4e542'), left=list(map(lambda do, gn, bn, grn, yn, on, vn, n, vs, gs, bs,
                                                                           grs: do + gn + bn + grn + yn + on + vn + n + vs + gs + bs + grs,
                                                                    dop_one, green_netto, blue_netto, gray_netto,
                                                                    yellow_netto, orange_netto, value_netto, names,
                                                                    value_ss, green_ss, blue_ss, gray_ss)))
                barh(y, orange_ss, color=('#ff9191'), left=list(map(
                    lambda do, gn, bn, grn, yn, on, vn, n, vs, gs, bs, grs,
                           ys: do + gn + bn + grn + yn + on + vn + n + vs + gs + bs + grs + ys, dop_one, green_netto,
                    blue_netto, gray_netto, yellow_netto, orange_netto, value_netto, names, value_ss, green_ss, blue_ss,
                    gray_ss, yellow_ss)))
                barh(y, dop_two, color='white', alpha=0, left=list(map(
                    lambda do, gn, bn, grn, yn, on, vn, n, vs, gs, bs, grs, os,
                           ys: do + gn + bn + grn + yn + on + vn + n + vs + gs + bs + grs + ys + os, dop_one,
                    green_netto, blue_netto, gray_netto, yellow_netto, orange_netto, value_netto, names, value_ss,
                    green_ss, blue_ss, gray_ss, yellow_ss, orange_ss)))

                # Подписи --------------------------------------------------------------------------------------
                ## суммы нетто
                ots = maxnetto + max(df.NettoS) / 1.2
                for i, ii in enumerate(y):
                    if i == len(y) - 1:
                        break
                    r = y[i]
                    Netto_label = '{:,}'.format(trunc(df.iloc[i][NettoSum].values[0])).replace(',', ' ')
                    plttext(ots, r, Netto_label, verticalalignment='center', horizontalalignment='right', **plot_font)

                ## суммы cc
                ots = maxnetto + max(df.NettoS) + max(df.IndeX) + max(df.SSS) / 1.2
                for i, ii in enumerate(y):
                    if i == len(y) - 1:
                        break
                    r = y[i]
                    SS_label = '{:,}'.format(trunc(df.iloc[i][SSSum].values[0])).replace(',', ' ')
                    plttext(ots, r, SS_label, verticalalignment='center', horizontalalignment='right', **plot_font)
                ## Статьи
                ots = maxnetto + max(df.NettoS) + max(df.IndeX) / 2
                for i, ii in enumerate(y):
                    if i == len(y) - 1:
                        break
                    r = y[i]
                    Index_label = df.iloc[i]['index'].values[0]
                    plttext(ots, r, Index_label, verticalalignment='center', horizontalalignment='center', **plot_font)

                ## Шапка
                plttext((maxnetto + maxnetto * 2 * 0.1) / 2, len(y) - 1, 'Балансовая стоимость',
                        verticalalignment='center', horizontalalignment='center', **plot_font_bold)
                plttext((maxnetto + maxnetto * 2 * 0.1) + maxnetto * 2 * 0.3 / 2, len(y) - 1, 'Наименование статьи',
                        verticalalignment='center', horizontalalignment='center', **plot_font_bold)
                plttext((maxnetto + maxnetto * 2 * 0.1) + maxnetto * 2 * 0.3 + (maxnetto + maxnetto * 2 * 0.1) / 2,
                        len(y) - 1, 'Справедливая стоимость', verticalalignment='center', horizontalalignment='center',
                        **plot_font_bold)
                plttext(0, len(y) - 1, 'в ' + izmerenie, verticalalignment='bottom', horizontalalignment='left',
                        **plot_font_bold)
                ## Итоги
                plttext(maxnetto + max(df.NettoS) / 1.2, -1, '{:,}'.format(trunc(SumNetto)).replace(',', ' '),
                        verticalalignment='center', horizontalalignment='right', **plot_font_bold)
                plttext((maxnetto + maxnetto * 2 * 0.1) + maxnetto * 2 * 0.3 / 2, -1, 'ИТОГО',
                        verticalalignment='center', horizontalalignment='center', **plot_font_bold)
                plttext(maxnetto + max(df.NettoS) + max(df.IndeX) + max(df.SSS) / 1.2, -1,
                        '{:,}'.format(trunc(SumSS)).replace(',', ' '), verticalalignment='center',
                        horizontalalignment='right', **plot_font_bold)

                # Подписи диаграмм
                if App.get_running_app().root.ids.xp.state == 'down':
                    ## столбец с суммами нетто
                    NS = squeeze(array(df[NettoSum]))
                    ## столбец с суммами сс
                    SSS = squeeze(array(df[SSSum]))
                    ## Деление
                    Dgn = green_netto[:-1] / NS[:-1]
                    Dbn = blue_netto[:-1] / NS[:-1]
                    Dgrn = gray_netto[:-1] / NS[:-1]
                    Dyn = yellow_netto[:-1] / NS[:-1]
                    Don = orange_netto[:-1] / NS[:-1]
                    # np.nan_to_num(green_ss[:-1] / SSS[:-1]) # более короткий способ деления на ноль, но выдает красные предупреждения
                    Dgs = divide(green_ss[:-1], SSS[:-1], out=zeros_like(green_ss[:-1]), where=SSS[:-1] != 0)
                    Dbs = divide(blue_ss[:-1], SSS[:-1], out=zeros_like(blue_ss[:-1]), where=SSS[:-1] != 0)
                    # Dys = np.divide(yellow_ss[:-1], SSS[:-1], out=np.zeros_like(yellow_ss[:-1]), where=SSS[:-1]!=0)
                    Dgrs = divide(gray_ss[:-1], SSS[:-1], out=zeros_like(gray_ss[:-1]), where=SSS[:-1] != 0)
                    ssots = maxnetto + max(df.NettoS) + max(df.IndeX) + max(
                        df.SSS)  # отступ для подписей СС bar (чтобы не вбивать много раз)
                    for i, j in enumerate(y):
                        if i == len(y):
                            break
                        if green_netto[i] / maxnetto >= 0.08:  # продажа
                            ots = dop_one[i] + green_netto[i] / 2
                            r = y[i]
                            if Dgn[i] != 1:
                                label = '{:,}'.format(trunc(green_netto[i])).replace(',',
                                                                                     ' ') + '\n(' + "{0:.0%}".format(
                                    Dgn[i]) + ')'
                            else:
                                label = '{:,}'.format(trunc(green_netto[i])).replace(',', ' ')
                            plttext(ots, r, label, verticalalignment='center', horizontalalignment='center',
                                    **plot_font_labl, fontsize=12)
                        else:
                            if green_netto[i] / maxnetto >= 0.05:
                                ots = dop_one[i] + green_netto[i] / 2
                                r = y[i]
                                if Dgn[i] != 1:
                                    label = "{0:.0%}".format(Dgn[i])
                                    plttext(ots, r, label, verticalalignment='center', horizontalalignment='center',
                                            **plot_font_labl)
                        if blue_netto[i] / maxnetto >= 0.08:  # до погашения
                            ots = dop_one[i] + green_netto[i] + blue_netto[i] / 2
                            r = y[i]
                            if Dbn[i] != 1:
                                label = '{:,}'.format(trunc(blue_netto[i])).replace(',',
                                                                                    ' ') + '\n(' + "{0:.0%}".format(
                                    Dbn[i]) + ')'
                            else:
                                label = '{:,}'.format(trunc(blue_netto[i])).replace(',', ' ')
                            plttext(ots, r, label, verticalalignment='center', horizontalalignment='center',
                                    **plot_font_labl, fontsize=12)
                        else:
                            if blue_netto[i] / maxnetto >= 0.05:
                                ots = dop_one[i] + green_netto[i] + blue_netto[i] / 2
                                r = y[i]
                                if Dbn[i] != 1:
                                    label = "{0:.0%}".format(Dbn[i])
                                    plttext(ots, r, label, verticalalignment='center', horizontalalignment='center',
                                            **plot_font_labl)
                        if gray_netto[i] / maxnetto >= 0.08:  # ТДИ
                            ots = dop_one[i] + green_netto[i] + blue_netto[i] + gray_netto[i] / 2
                            r = y[i]
                            if Dgrn[i] != 1:
                                label = '{:,}'.format(trunc(gray_netto[i])).replace(',',
                                                                                    ' ') + '\n(' + "{0:.0%}".format(
                                    Dgrn[i]) + ')'
                            else:
                                label = '{:,}'.format(trunc(gray_netto[i])).replace(',', ' ')
                            plttext(ots, r, label, verticalalignment='center', horizontalalignment='center',
                                    **plot_font_labl, fontsize=12)
                        else:
                            if gray_netto[i] / maxnetto >= 0.05:
                                ots = dop_one[i] + green_netto[i] + blue_netto[i] + gray_netto[i] / 2
                                r = y[i]
                                if Dgrn[i] != 1:
                                    label = "{0:.0%}".format(Dgrn[i])
                                    plttext(ots, r, label, verticalalignment='center', horizontalalignment='center',
                                            **plot_font_labl)
                        if yellow_netto[i] / maxnetto >= 0.08:  # техника
                            ots = dop_one[i] + green_netto[i] + blue_netto[i] + gray_netto[i] + yellow_netto[i] / 2
                            r = y[i]
                            if Dyn[i] != 1:
                                label = '{:,}'.format(trunc(yellow_netto[i])).replace(',',
                                                                                      ' ') + '\n(' + "{0:.0%}".format(
                                    Dyn[i]) + ')'
                            else:
                                label = '{:,}'.format(trunc(yellow_netto[i])).replace(',', ' ')
                            plttext(ots, r, label, verticalalignment='center', horizontalalignment='center',
                                    **plot_font_labl, fontsize=12)
                        else:
                            if yellow_netto[i] / maxnetto >= 0.05:
                                ots = dop_one[i] + green_netto[i] + blue_netto[i] + gray_netto[i] + yellow_netto[i] / 2
                                r = y[i]
                                if Dyn[i] != 1:
                                    label = "{0:.0%}".format(Dyn[i])
                                    plttext(ots, r, label, verticalalignment='center', horizontalalignment='center',
                                            **plot_font_labl)
                        if orange_netto[i] / maxnetto >= 0.08:  # списание
                            ots = dop_one[i] + green_netto[i] + blue_netto[i] + gray_netto[i] + yellow_netto[i] + \
                                  orange_netto[i] / 2
                            r = y[i]
                            if Don[i] != 1:
                                label = '{:,}'.format(trunc(orange_netto[i])).replace(',',
                                                                                      ' ') + '\n(' + "{0:.0%}".format(
                                    Don[i]) + ')'
                            else:
                                label = '{:,}'.format(trunc(orange_netto[i])).replace(',', ' ')
                            plttext(ots, r, label, verticalalignment='center', horizontalalignment='center',
                                    **plot_font_labl, fontsize=12)
                        else:
                            if orange_netto[i] / maxnetto >= 0.05:
                                ots = dop_one[i] + green_netto[i] + blue_netto[i] + gray_netto[i] + yellow_netto[i] + \
                                      orange_netto[i] / 2
                                r = y[i]
                                if Don[i] != 1:
                                    label = "{0:.0%}".format(Don[i])
                                    plttext(ots, r, label, verticalalignment='center', horizontalalignment='center',
                                            **plot_font_labl)
                        if green_ss[i] / maxnetto >= 0.08:  # продажа_СС
                            ots = ssots + green_ss[i] / 2
                            r = y[i]
                            if Dgs[i] != 1:
                                label = '{:,}'.format(trunc(green_ss[i])).replace(',', ' ') + '\n(' + "{0:.0%}".format(
                                    Dgs[i]) + ')'
                            else:
                                label = '{:,}'.format(trunc(green_ss[i])).replace(',', ' ')
                            plttext(ots, r, label, verticalalignment='center', horizontalalignment='center',
                                    **plot_font_labl, fontsize=12)
                        else:
                            if green_ss[i] / maxnetto >= 0.05:
                                ots = ssots + green_ss[i] / 2
                                r = y[i]
                                if Dgs[i] != 1:
                                    label = "{0:.0%}".format(Dgs[i])
                                    plttext(ots, r, label, verticalalignment='center', horizontalalignment='center',
                                            **plot_font_labl)
                        if blue_ss[i] / maxnetto >= 0.08:  # до погашения_СС
                            ots = ssots + green_ss[i] + blue_ss[i] / 2
                            r = y[i]
                            if Dbs[i] != 1:
                                label = '{:,}'.format(trunc(blue_ss[i])).replace(',', ' ') + '\n(' + "{0:.0%}".format(
                                    Dbs[i]) + ')'
                            else:
                                label = '{:,}'.format(trunc(blue_ss[i])).replace(',', ' ')
                            plttext(ots, r, label, verticalalignment='center', horizontalalignment='center',
                                    **plot_font_labl, fontsize=12)
                        else:
                            if blue_ss[i] / maxnetto >= 0.05:
                                ots = ssots + green_ss[i] + blue_ss[i] / 2
                                r = y[i]
                                if Dbs[i] != 1:
                                    label = "{0:.0%}".format(Dbs[i])
                                    plttext(ots, r, label, verticalalignment='center', horizontalalignment='center',
                                            **plot_font_labl)
                        if gray_ss[i] / maxnetto >= 0.08:  # ТДИ_СС
                            ots = ssots + green_ss[i] + blue_ss[i] + gray_ss[i] / 2
                            r = y[i]
                            if Dgrs[i] != 1:
                                label = '{:,}'.format(trunc(gray_ss[i])).replace(',', ' ') + '\n(' + "{0:.0%}".format(
                                    Dgrs[i]) + ')'
                            else:
                                label = '{:,}'.format(trunc(gray_ss[i])).replace(',', ' ')
                            plttext(ots, r, label, verticalalignment='center', horizontalalignment='center',
                                    **plot_font_labl, fontsize=12)
                        else:
                            if gray_ss[i] / maxnetto >= 0.05:
                                ots = ssots + green_ss[i] + blue_ss[i] + gray_ss[i] / 2
                                r = y[i]
                                if Dgrs[i] != 1:
                                    label = "{0:.0%}".format(Dgrs[i])
                                    plttext(ots, r, label, verticalalignment='center', horizontalalignment='center',
                                            **plot_font_labl)
                        if yellow_ss[i] / maxnetto >= 0.08:  # Техника_СС
                            ots = ssots + green_ss[i] + blue_ss[i] + gray_ss[i] + yellow_ss[i] / 2
                            r = y[i]
                            label = '{:,}'.format(trunc(yellow_ss[i])).replace(',', ' ')
                            plttext(ots, r, label, verticalalignment='center', horizontalalignment='center',
                                    **plot_font_labl, fontsize=12)
                        if orange_ss[i] / maxnetto >= 0.07:  # списание_СС
                            ots = ssots + green_ss[i] + blue_ss[i] + gray_ss[i] + yellow_ss[i] + orange_ss[i] / 2
                            r = y[i]
                            label = '{:,}'.format(trunc(orange_ss[i])).replace(',', ' ')
                            plttext(ots, r, label, verticalalignment='center', horizontalalignment='center',
                                    **plot_font_labl, fontsize=12)
                # ----------------------------------------------------------------------------------------------

                # сетка
                plot((maxnetto + maxnetto * 2 * 0.1, maxnetto + maxnetto * 2 * 0.1), (y[0] - 0.5, y[-1] - 0.5),
                     alpha=0.3, color=('#02003d'), linestyle='--', linewidth=1)
                plot((maxnetto + maxnetto * 2 * 0.1 + maxnetto * 2 * 0.3,
                      maxnetto + maxnetto * 2 * 0.1 + maxnetto * 2 * 0.3), (y[0] - 0.5, y[-1] - 0.5), alpha=0.3,
                     color=('#02003d'), linestyle='--', linewidth=1)
                for i in y:
                    if i == 0:  # Линия итогов
                        plot((0, maxnetto + maxnetto * 2 * 0.1 + maxnetto * 2 * 0.3 + maxnetto * 2 * 0.1 + maxnetto),
                             (i - 0.5, i - 0.5), alpha=1, color=('#02003d'), linewidth=2)
                    if i == y[-1]:  # Линия шапки
                        plot((0, maxnetto + maxnetto * 2 * 0.1 + maxnetto * 2 * 0.3 + maxnetto * 2 * 0.1 + maxnetto),
                             (i - 0.5, i - 0.5), alpha=1, color=('#02003d'), linewidth=2)
                    else:
                        plot((0, maxnetto + maxnetto * 2 * 0.1 + maxnetto * 2 * 0.3 + maxnetto * 2 * 0.1 + maxnetto),
                             (i - 0.5, i - 0.5), alpha=0.3, color=('#02003d'), linestyle='--', linewidth=1)

                # Легенда
                legend(bbox_to_anchor=(0, 0), loc='center left', ncol=3, framealpha=0.0,
                       prop={'size': 12})  # prop={"family":"Arial", 'stretch' : 'condensed'}

            if n == 2:
                ## Проставляем значения столбцов, _net - обесценение
                h1 = array(DF.iloc[:, 0])
                h2 = array(DF.iloc[:, 1])
                h3 = array(DF.iloc[:, 2])
                names = nan_to_num(array(DF.iloc[:, len(spisok) + 5]))
                value_netto = nan_to_num(array(DF.iloc[:, len(spisok) + 4])) * 2
                value_ss = nan_to_num(array(DF.iloc[:, len(spisok) + 6])) * 2
                dop_one = names + value_netto + value_ss
                # 24 13 13 50
                try:
                    green_ss = nan_to_num(array(DF.iloc[:, green_ss_col])).astype('float')  # /
                except:
                    green_ss = zeros(len(DF.iloc[:, 4]))
                try:
                    blue_ss = nan_to_num(array(DF.iloc[:, blue_ss_col])).astype('float')
                except:
                    blue_ss = zeros(len(DF.iloc[:, 4]))
                try:
                    gray_ss = nan_to_num(array(DF.iloc[:, gray_ss_col])).astype('float')
                except:
                    gray_ss = zeros(len(DF.iloc[:, 4]))
                try:
                    yellow_ss = nan_to_num(array(DF.iloc[:, yellow_ss_col])).astype('float')
                except:
                    yellow_ss = zeros(len(DF.iloc[:, 4]))
                try:
                    orange_ss = nan_to_num(array(DF.iloc[:, orange_ss_col])).astype('float')
                except:
                    orange_ss = zeros(len(DF.iloc[:, 4]))
                try:
                    green_net = nan_to_num(DF.iloc[:, green_netto_col]).astype('float') - green_ss
                except:
                    green_net = zeros(len(DF.iloc[:, 4]))
                try:
                    blue_net = nan_to_num(array(DF.iloc[:, blue_netto_col])).astype('float') - blue_ss
                except:
                    blue_net = zeros(len(DF.iloc[:, 4]))
                try:
                    gray_net = nan_to_num(array(DF.iloc[:, gray_netto_col])).astype('float') - gray_ss
                except:
                    gray_net = zeros(len(DF.iloc[:, 4]))
                try:
                    yellow_net = nan_to_num(array(DF.iloc[:, yellow_netto_col])).astype('float') - yellow_ss
                except:
                    yellow_net = zeros(len(DF.iloc[:, 4]))
                try:
                    orange_net = nan_to_num(array(DF.iloc[:, orange_netto_col])).astype('float') - orange_ss
                except:
                    orange_net = zeros(len(DF.iloc[:, 4]))
                # отрисовка диаграмм
                barh(y, h1, color=('#02003d'), height=1.2, alpha=0)
                barh(y, h2, color=('#02003d'), height=1.2, left=h1, alpha=0)
                barh(y, h3, color=('#02003d'), height=1.2, left=list(map(lambda h, hh: h + hh, h1, h2)), alpha=0)
                barh(y, names, color='white', alpha=0)
                barh(y, value_netto, color='white', alpha=0, left=names)
                if App.get_running_app().root.ids.xo.state == 'down':
                    dop_one = dop_one - value_ss
                else:
                    barh(y, value_ss, color='white', alpha=0, left=list(map(lambda n, vn: n + vn, names, value_netto)))
                barh(y, green_ss, color=('#8cba88'), label=greenlabel, left=dop_one)
                barh(y, green_net, color=('#8cba88'), left=list(map(lambda d, gs: d + gs, dop_one, green_ss)),
                     hatch='//')
                barh(y, blue_ss, color=('#79a5d4'), label=bluelabel,
                     left=list(map(lambda d, gs, gn: d + gs + gn, dop_one, green_ss, green_net)))
                barh(y, blue_net, color=('#79a5d4'),
                     left=list(map(lambda d, gs, gn, bs: d + gs + gn + bs, dop_one, green_ss, green_net, blue_ss)),
                     hatch='//')
                barh(y, gray_ss, color=('#878787'), label=graylabel, left=list(
                    map(lambda d, gs, gn, bs, bn: d + gs + gn + bs + bn, dop_one, green_ss, green_net, blue_ss,
                        blue_net)))
                barh(y, gray_net, color=('#878787'), left=list(
                    map(lambda d, gs, gn, bs, bn, grs: d + gs + gn + bs + bn + grs, dop_one, green_ss, green_net,
                        blue_ss, blue_net, gray_ss)), hatch='//')
                barh(y, yellow_ss, color=('#f4e542'), label=yelowlabel, left=list(
                    map(lambda d, gs, gn, bs, bn, grs, grn: d + gs + gn + bs + bn + grs + grn, dop_one, green_ss,
                        green_net, blue_ss, blue_net, gray_ss, gray_net)))
                barh(y, yellow_net, color=('#f4e542'), left=list(
                    map(lambda d, gs, gn, bs, bn, grs, grn, ys: d + gs + gn + bs + bn + grs + grn + ys, dop_one,
                        green_ss, green_net, blue_ss, blue_net, gray_ss, gray_net, yellow_ss)), hatch='//')
                barh(y, orange_ss, color=('#ff9191'), label=orangelabel, left=list(
                    map(lambda d, gs, gn, bs, bn, grs, grn, ys, yn: d + gs + gn + bs + bn + grs + grn + ys + yn,
                        dop_one, green_ss, green_net, blue_ss, blue_net, gray_ss, gray_net, yellow_ss, yellow_net)))
                barh(y, orange_net, color=('#ff9191'), left=list(map(
                    lambda d, gs, gn, bs, bn, grs, grn, ys, yn, os: d + gs + gn + bs + bn + grs + grn + ys + yn + os,
                    dop_one, green_ss, green_net, blue_ss, blue_net, gray_ss, gray_net, yellow_ss, yellow_net,
                    orange_ss)), hatch='//')
                if App.get_running_app().root.ids.xo.state == 'down':
                    barh(y, value_ss, color='white', alpha=0, left=list(map(
                        lambda d, gs, gn, bs, bn, grs, grn, ys, yn, os,
                               on: d + gs + gn + bs + bn + grs + grn + ys + yn + os + on, dop_one, green_ss, green_net,
                        blue_ss, blue_net, gray_ss, gray_net, yellow_ss, yellow_net, orange_ss, orange_net)))
                # barh(0, 0, color='white', hatch='//', label='Обесценение')
                # Подписи --------------------------------------------------------------------------------------
                ## Статьи
                for i, ii in enumerate(y):
                    if i == len(y) - 1:
                        break
                    r = y[i]
                    Index_label = df.iloc[i]['index'].values[0]
                    plttext(0, r, Index_label, verticalalignment='center', horizontalalignment='left', **plot_font)

                ## суммы нетто
                ots = max(df.IndeX) + max(df.NettoS) * 2 / 1.2
                for i, ii in enumerate(y):
                    if i == len(y) - 1:
                        break
                    r = y[i]
                    Netto_label = '{:,}'.format(trunc(df.iloc[i][NettoSum].values[0])).replace(',', ' ')
                    plttext(ots, r, Netto_label, verticalalignment='center', horizontalalignment='right', **plot_font)

                ## суммы cc
                if App.get_running_app().root.ids.xo.state == 'down':
                    ots = max(df.IndeX) + max(df.NettoS) * 2 + maxnetto + max(df.SSS) * 2 / 1.2
                else:
                    ots = max(df.IndeX) + max(df.NettoS) * 2 + max(df.SSS) * 2 / 1.2
                for i, ii in enumerate(y):
                    if i == len(y) - 1:
                        break
                    r = y[i]
                    SS_label = '{:,}'.format(trunc(df.iloc[i][SSSum].values[0])).replace(',', ' ')
                    plttext(ots, r, SS_label, verticalalignment='center', horizontalalignment='right', **plot_font)

                ## Шапка
                plttext(0, len(y) - 1, 'Наименование статьи, \n' + izmerenie, verticalalignment='center',
                        horizontalalignment='left', **plot_font_bold)
                plttext(max(df.IndeX) + max(df.NettoS) * 2 / 2, len(y) - 1, 'Балансовая\nстоимость',
                        verticalalignment='center', horizontalalignment='center', **plot_font_bold)
                if App.get_running_app().root.ids.xo.state == 'down':
                    plttext(max(df.IndeX) + max(df.NettoS) * 2, len(y) - 1, 'Распределение активов по группам',
                            verticalalignment='center', horizontalalignment='left', **plot_font_bold)
                    plttext(max(df.IndeX) + max(df.NettoS) * 2 + maxnetto + max(df.NettoS) * 2 / 2, len(y) - 1,
                            'Справедливая\nстоимость', verticalalignment='center', horizontalalignment='center',
                            **plot_font_bold)
                else:
                    plttext(max(df.IndeX) + max(df.NettoS) * 2 + max(df.SSS) * 2 / 2, len(y) - 1,
                            'Справедливая\nстоимость', verticalalignment='center', horizontalalignment='center',
                            **plot_font_bold)
                    plttext(max(df.IndeX) + max(df.NettoS) * 2 + max(df.SSS) * 2, len(y) - 1,
                            'Распределение активов по группам', verticalalignment='center', horizontalalignment='left',
                            **plot_font_bold)

                ## Итоги
                plttext(0, -1, 'ИТОГО', verticalalignment='center', horizontalalignment='left', **plot_font_bold)
                plttext(max(df.IndeX) + max(df.NettoS) * 2 / 1.2, -1, '{:,}'.format(trunc(SumNetto)).replace(',', ' '),
                        verticalalignment='center', horizontalalignment='right', **plot_font_bold)
                if App.get_running_app().root.ids.xo.state == 'down':
                    plttext(max(df.IndeX) + max(df.NettoS) * 2 + maxnetto + max(df.SSS) * 2 / 1.2, -1,
                            '{:,}'.format(trunc(SumSS)).replace(',', ' '), verticalalignment='center',
                            horizontalalignment='right', **plot_font_bold)
                else:
                    plttext(max(df.IndeX) + max(df.NettoS) * 2 + max(df.SSS) * 2 / 1.2, -1,
                            '{:,}'.format(trunc(SumSS)).replace(',', ' '), verticalalignment='center',
                            horizontalalignment='right', **plot_font_bold)
                # Подписи диаграмм
                ## столбец с суммами сс
                if App.get_running_app().root.ids.xp.state == 'down':
                    SSS = squeeze(array(df[SSSum]))
                    if App.get_running_app().root.ids.xo.state == 'down':
                        ssots = max(df.IndeX) + max(
                            df.NettoS) * 2  # отступ для подписей СС bar (чтобы не вбивать много раз)
                    else:
                        ssots = max(df.IndeX) + max(df.NettoS) * 2 + max(
                            df.SSS) * 2  # отступ для подписей СС bar (чтобы не вбивать много раз)
                    for i, j in enumerate(y):
                        if i == len(y):
                            break
                        if green_ss[i] / maxnetto >= 0.08:  # продажа_СС
                            ots = ssots + green_ss[i] / 2
                            r = y[i]
                            label = '{:,}'.format(trunc(green_ss[i])).replace(',', ' ')
                            plttext(ots, r, label, verticalalignment='center', horizontalalignment='center',
                                    **plot_font_labl, fontsize=14)
                        if blue_ss[i] / maxnetto >= 0.08:  # до погашения_СС
                            ots = ssots + green_ss[i] + green_net[i] + blue_ss[i] / 2
                            r = y[i]
                            label = '{:,}'.format(trunc(blue_ss[i])).replace(',', ' ')
                            plttext(ots, r, label, verticalalignment='center', horizontalalignment='center',
                                    **plot_font_labl, fontsize=14)
                        if gray_ss[i] / maxnetto >= 0.08:  # ТДИ_СС
                            ots = ssots + green_ss[i] + green_net[i] + blue_ss[i] + blue_net[i] + gray_ss[i] / 2
                            r = y[i]
                            label = '{:,}'.format(trunc(gray_ss[i])).replace(',', ' ')
                            plttext(ots, r, label, verticalalignment='center', horizontalalignment='center',
                                    **plot_font_labl, fontsize=14)
                # сетка
                for i in y:
                    if i == 0:  # Линия итогов
                        plot((0, maxnetto + maxnetto * 2 * 0.1 + maxnetto * 2 * 0.3 + maxnetto * 2 * 0.1 + maxnetto),
                             (i - 0.5, i - 0.5), alpha=1, color=('#02003d'), linewidth=2)
                    if i == y[-1]:  # Линия шапки
                        plot((0, maxnetto + maxnetto * 2 * 0.1 + maxnetto * 2 * 0.3 + maxnetto * 2 * 0.1 + maxnetto),
                             (i - 0.5, i - 0.5), alpha=1, color=('#02003d'), linewidth=2)
                    else:
                        plot((0, maxnetto + maxnetto * 2 * 0.1 + maxnetto * 2 * 0.3 + maxnetto * 2 * 0.1 + maxnetto),
                             (i - 0.5, i - 0.5), alpha=0.3, color=('#02003d'), linestyle='--', linewidth=1)

                # Легенда
                if App.get_running_app().root.ids.xo.state == 'down':
                    legend(bbox_to_anchor=(0.308, 0), loc='center left', ncol=3, framealpha=0.0, prop={'size': 12})
                else:
                    legend(bbox_to_anchor=(1, 0), loc='center right', ncol=3, framealpha=0.0,
                           prop={'size': 12})  # prop={"family":"Arial", 'stretch' : 'condensed'}

            if n == 3:
                ## Проставляем значения столбцов, _net - обесценение
                h1 = array(DF.iloc[:, 0])
                h2 = array(DF.iloc[:, 1])
                h3 = array(DF.iloc[:, 2])
                names = nan_to_num(array(DF.iloc[:, len(spisok) + 5]))
                value_netto = nan_to_num(array(DF.iloc[:, len(spisok) + 4])) * 2
                value_ss = nan_to_num(array(DF.iloc[:, len(spisok) + 6])) * 2
                dop_one = names + value_netto + value_ss
                try:
                    green_ss = nan_to_num(array(DF.iloc[:, green_ss_col])).astype('float')
                except:
                    green_ss = zeros(len(DF.iloc[:, 4]))
                try:
                    blue_ss = nan_to_num(array(DF.iloc[:, blue_ss_col])).astype('float')
                except:
                    blue_ss = zeros(len(DF.iloc[:, 4]))
                try:
                    gray_ss = nan_to_num(array(DF.iloc[:, gray_ss_col])).astype('float')
                except:
                    gray_ss = zeros(len(DF.iloc[:, 4]))
                try:
                    yellow_ss = nan_to_num(array(DF.iloc[:, yellow_ss_col])).astype('float')
                except:
                    yellow_ss = zeros(len(DF.iloc[:, 4]))
                try:
                    orange_ss = nan_to_num(array(DF.iloc[:, orange_ss_col])).astype('float')
                except:
                    orange_ss = zeros(len(DF.iloc[:, 4]))
                try:
                    green_net = nan_to_num(DF.iloc[:, green_netto_col]).astype('float') - green_ss
                    # Считаем допы для отсутсупов
                    maxgn = max(nan_to_num(DF.iloc[:, green_netto_col]).astype('float')) - nan_to_num(
                        DF.iloc[:, green_netto_col]).astype('float')
                except:
                    green_net = zeros(len(DF.iloc[:, 4]))
                    maxgn = green_net
                try:
                    blue_net = nan_to_num(array(DF.iloc[:, blue_netto_col])).astype('float') - blue_ss
                    maxbn = max(nan_to_num(array(DF.iloc[:, blue_netto_col])).astype('float')) - nan_to_num(
                        array(DF.iloc[:, blue_netto_col])).astype('float')
                except:
                    blue_net = zeros(len(DF.iloc[:, 4]))
                    maxbn = blue_net
                try:
                    gray_net = nan_to_num(array(DF.iloc[:, gray_netto_col])).astype('float') - gray_ss
                    maxgrn = max(nan_to_num(array(DF.iloc[:, gray_netto_col])).astype('float')) - nan_to_num(
                        array(DF.iloc[:, gray_netto_col])).astype('float')
                except:
                    gray_net = zeros(len(DF.iloc[:, 4]))
                    maxgrn = gray_net
                try:
                    yellow_net = nan_to_num(array(DF.iloc[:, yellow_netto_col])).astype('float') - yellow_ss
                    maxyn = max(nan_to_num(array(DF.iloc[:, yellow_netto_col])).astype('float')) - nan_to_num(
                        array(DF.iloc[:, yellow_netto_col])).astype('float')
                except:
                    yellow_net = zeros(len(DF.iloc[:, 4]))
                    maxyn = yellow_net
                try:
                    orange_net = nan_to_num(array(DF.iloc[:, orange_netto_col])).astype('float') - orange_ss
                except:
                    orange_net = zeros(len(DF.iloc[:, 4]))

                # отрисовка диаграмм
                barh(y, h1, color=('#02003d'), height=1.2, alpha=0)
                barh(y, h2, color=('#02003d'), height=1.2, left=h1, alpha=0)
                barh(y, h3, color=('#02003d'), height=1.2, left=list(map(lambda h, hh: h + hh, h1, h2)), alpha=0)
                barh(y, names, color='white', alpha=0)
                barh(y, value_netto, color='white', alpha=0, left=names)
                barh(y, value_ss, color='white', alpha=0, left=list(map(lambda n, vn: n + vn, names, value_netto)))
                barh(y, green_ss, color=('#8cba88'), label=greenlabel, left=dop_one)
                barh(y, green_net, color=('#8cba88'), left=list(map(lambda d, gs: d + gs, dop_one, green_ss)),
                     hatch='//')
                barh(y, maxgn, color=('#8cba88'),
                     left=list(map(lambda d, gs, gn: d + gs + gn, dop_one, green_ss, green_net)), alpha=0)
                barh(y, blue_ss, color=('#79a5d4'), label=bluelabel,
                     left=list(map(lambda d, gs, gn, mgn: d + gs + gn + mgn, dop_one, green_ss, green_net, maxgn)))
                barh(y, blue_net, color=('#79a5d4'), left=list(
                    map(lambda d, gs, gn, mgn, bs: d + gs + gn + mgn + bs, dop_one, green_ss, green_net, maxgn,
                        blue_ss)), hatch='//')
                barh(y, maxbn, color=('#878787'), left=list(
                    map(lambda d, gs, gn, mgn, bs, bn: d + gs + gn + mgn + bs + bn, dop_one, green_ss, green_net, maxgn,
                        blue_ss, blue_net)), alpha=0)
                barh(y, gray_ss, color=('#878787'), label=graylabel, left=list(
                    map(lambda d, gs, gn, mgn, bs, bn, mbn: d + gs + gn + mgn + bs + bn + mbn, dop_one, green_ss,
                        green_net, maxgn, blue_ss, blue_net, maxbn)))
                barh(y, gray_net, color=('#878787'), left=list(
                    map(lambda d, gs, gn, mgn, bs, bn, mbn, grs: d + gs + gn + mgn + bs + bn + mbn + grs, dop_one,
                        green_ss, green_net, maxgn, blue_ss, blue_net, maxbn, gray_ss)), hatch='//')
                barh(y, maxgrn, color=('#f4e542'), alpha=0, left=list(
                    map(lambda d, gs, gn, mgn, bs, bn, mbn, grs, grn: d + gs + gn + mgn + bs + bn + mbn + grs + grn,
                        dop_one, green_ss, green_net, maxgn, blue_ss, blue_net, maxbn, gray_ss, gray_net)))
                barh(y, yellow_ss, color=('#f4e542'), label=yelowlabel, left=list(map(
                    lambda d, gs, gn, mgn, bs, bn, mbn, grs, grn,
                           mgrn: d + gs + gn + mgn + bs + bn + mbn + grs + grn + mgrn, dop_one, green_ss, green_net,
                    maxgn, blue_ss, blue_net, maxbn, gray_ss, gray_net, maxgrn)))
                barh(y, yellow_net, color=('#f4e542'), left=list(map(lambda d, gs, gn, mgn, bs, bn, mbn, grs, grn, mgrn,
                                                                            ys: d + gs + gn + mgn + bs + bn + mbn + grs + grn + mgrn + ys,
                                                                     dop_one, green_ss, green_net, maxgn, blue_ss,
                                                                     blue_net, maxbn, gray_ss, gray_net, maxgrn,
                                                                     yellow_ss)), hatch='//')
                barh(y, maxyn, color=('#ff9191'), alpha=0, left=list(map(
                    lambda d, gs, gn, mgn, bs, bn, mbn, grs, grn, mgrn, ys,
                           yn: d + gs + gn + mgn + bs + bn + mbn + grs + grn + mgrn + ys + yn, dop_one, green_ss,
                    green_net, maxgn, blue_ss, blue_net, maxbn, gray_ss, gray_net, maxgrn, yellow_ss, yellow_net)))
                barh(y, orange_ss, color=('#ff9191'), label=orangelabel, left=list(map(
                    lambda d, gs, gn, mgn, bs, bn, mbn, grs, grn, mgrn, ys, yn,
                           myn: d + gs + gn + mgn + bs + bn + mbn + grs + grn + mgrn + ys + yn + myn, dop_one, green_ss,
                    green_net, maxgn, blue_ss, blue_net, maxbn, gray_ss, gray_net, maxgrn, yellow_ss, yellow_net,
                    maxyn)))
                barh(y, orange_net, color=('#ff9191'), left=list(map(
                    lambda d, gs, gn, mgn, bs, bn, mbn, grs, grn, mgrn, ys, yn, myn,
                           os: d + gs + gn + mgn + bs + bn + mbn + grs + grn + mgrn + ys + yn + myn + os, dop_one,
                    green_ss, green_net, maxgn, blue_ss, blue_net, maxbn, gray_ss, gray_net, maxgrn, yellow_ss,
                    yellow_net, maxyn, orange_ss)), hatch='//')
                barh(0, 0, color='white', hatch='//', label='Обесценение')
                # Подписи --------------------------------------------------------------------------------------
                ## Статьи
                for i, ii in enumerate(y):
                    if i == len(y) - 1:
                        break
                    r = y[i]
                    Index_label = df.iloc[i]['index'].values[0]
                    plttext(0, r, Index_label, verticalalignment='center', horizontalalignment='left', **plot_font)

                ## суммы нетто
                ots = max(df.IndeX) + max(df.NettoS) * 2 / 1.2
                for i, ii in enumerate(y):
                    if i == len(y) - 1:
                        break
                    r = y[i]
                    Netto_label = '{:,}'.format(trunc(df.iloc[i][NettoSum].values[0])).replace(',', ' ')
                    plttext(ots, r, Netto_label, verticalalignment='center', horizontalalignment='right', **plot_font)

                ## суммы cc
                ots = max(df.IndeX) + max(df.NettoS) * 2 + max(df.SSS) * 2 / 1.2
                for i, ii in enumerate(y):
                    if i == len(y) - 1:
                        break
                    r = y[i]
                    SS_label = '{:,}'.format(trunc(df.iloc[i][SSSum].values[0])).replace(',', ' ')
                    plttext(ots, r, SS_label, verticalalignment='center', horizontalalignment='right', **plot_font)

                ## Шапка
                plttext(0, len(y) - 1, 'Наименование статьи, \n' + izmerenie, verticalalignment='center',
                        horizontalalignment='left', **plot_font_bold)
                plttext(max(df.IndeX) + max(df.NettoS) * 2 / 2, len(y) - 1, 'Балансовая\nстоимость',
                        verticalalignment='center', horizontalalignment='center', **plot_font_bold)
                plttext(max(df.IndeX) + max(df.NettoS) * 2 + max(df.SSS) * 2 / 2, len(y) - 1, 'Справедливая\nстоимость',
                        verticalalignment='center', horizontalalignment='center', **plot_font_bold)
                plttext(max(df.IndeX) + max(df.NettoS) * 2 + max(df.SSS) * 2, len(y) - 1,
                        'Распределение активов по группам', verticalalignment='center', horizontalalignment='left',
                        **plot_font_bold)

                ## Итоги
                plttext(0, -1, 'ИТОГО', verticalalignment='center', horizontalalignment='left', **plot_font_bold)
                plttext(max(df.IndeX) + max(df.NettoS) * 2 / 1.2, -1, '{:,}'.format(trunc(SumNetto)).replace(',', ' '),
                        verticalalignment='center', horizontalalignment='right', **plot_font_bold)
                plttext(max(df.IndeX) + max(df.NettoS) * 2 + max(df.SSS) * 2 / 1.2, -1,
                        '{:,}'.format(trunc(SumSS)).replace(',', ' '), verticalalignment='center',
                        horizontalalignment='right', **plot_font_bold)
                # Подписи диаграмм
                ## столбец с суммами сс
                if App.get_running_app().root.ids.xp.state == 'down':
                    SSS = squeeze(array(df[SSSum]))
                    ssots = max(df.IndeX) + max(df.NettoS) * 2 + max(
                        df.SSS) * 2  # отступ для подписей СС bar (чтобы не вбивать много раз)
                    for i, j in enumerate(y):
                        if i == len(y):
                            break
                        if green_ss[i] / maxnetto >= 0.08:  # продажа_СС
                            ots = ssots + green_ss[i] / 2
                            r = y[i]
                            label = '{:,}'.format(trunc(green_ss[i])).replace(',', ' ')
                            plttext(ots, r, label, verticalalignment='center', horizontalalignment='center',
                                    **plot_font_labl, fontsize=14)
                        if blue_ss[i] / maxnetto >= 0.08:  # до погашения_СС
                            ots = ssots + green_ss[i] + green_net[i] + maxgn[i] + blue_ss[i] / 2
                            r = y[i]
                            label = '{:,}'.format(trunc(blue_ss[i])).replace(',', ' ')
                            plttext(ots, r, label, verticalalignment='center', horizontalalignment='center',
                                    **plot_font_labl, fontsize=14)
                        if gray_ss[i] / maxnetto >= 0.08:  # ТДИ_СС
                            ots = ssots + green_ss[i] + green_net[i] + maxgn[i] + blue_ss[i] + blue_net[i] + maxbn[i] + \
                                  gray_ss[i] / 2
                            r = y[i]
                            label = '{:,}'.format(trunc(gray_ss[i])).replace(',', ' ')
                            plttext(ots, r, label, verticalalignment='center', horizontalalignment='center',
                                    **plot_font_labl, fontsize=14)
                for i in y:
                    if i == 0:  # Линия итогов
                        plot((0, maxnetto + maxnetto * 2 * 0.1 + maxnetto * 2 * 0.3 + maxnetto * 2 * 0.1 + maxnetto),
                             (i - 0.5, i - 0.5), alpha=1, color=('#02003d'), linewidth=2)
                    if i == y[-1]:  # Линия шапки
                        plot((0, maxnetto + maxnetto * 2 * 0.1 + maxnetto * 2 * 0.3 + maxnetto * 2 * 0.1 + maxnetto),
                             (i - 0.5, i - 0.5), alpha=1, color=('#02003d'), linewidth=2)
                    else:
                        plot((0, maxnetto + maxnetto * 2 * 0.1 + maxnetto * 2 * 0.3 + maxnetto * 2 * 0.1 + maxnetto),
                             (i - 0.5, i - 0.5), alpha=0.3, color=('#02003d'), linestyle='--', linewidth=1)

                # Легенда
                legend(bbox_to_anchor=(1, 0), loc='center right', ncol=3, framealpha=0.0,
                       prop={'size': 12})  # prop={"family":"Arial", 'stretch' : 'condensed'}

            savefig(osexp('~/Desktop/Plot.png'), bbox_inches='tight', pad_inches=0)
            pltclose()
        except Exception as e:
            print(e)
            Popup().open()

    # ОСНОВНАЯ ФУНКЦИЯ ВЫПИСКАТОРА
    def Visualize(self):
        def Azbuka(self, x):
            Excel = {'A': 0, 'B': 1, 'C': 2, 'D': 3, 'E': 4, 'F': 5,
                     'G': 6, 'H': 7, 'I': 8, 'J': 9, 'K': 10, 'L': 11, 'M': 12,
                     'N': 13, 'O': 14, 'P': 15, 'Q': 16, 'R': 17, 'S': 18, 'T': 19,
                     'U': 20, 'V': 21, 'W': 22, 'X': 23, 'Y': 24, 'Z': 25, 'AA': 26,
                     'AB': 27, 'AC': 28, 'AD': 29, 'AE': 30, 'AF': 31, 'AG': 32, 'AH': 33,
                     'AI': 34, 'AJ': 35, 'AK': 36, 'AL': 37, 'AM': 38, 'AN': 39, 'AO': 40,
                     'AP': 41, 'AQ': 42, 'AR': 43, 'AS': 44, 'AT': 45, 'AU': 46, 'AV': 47,
                     'AW': 48, 'AX': 49, 'AY': 50, 'AZ': 51, 'BA': 52, 'BB': 53, 'BC': 54,
                     'BD': 55, 'BE': 56, 'BF': 57, 'BG': 58, 'BH': 59, 'BI': 60, 'BJ': 61,
                     'BK': 62, 'BL': 63, 'BM': 64, 'BN': 65, 'BO': 66, 'BP': 67, 'BQ': 68,
                     'BR': 69, 'BS': 70, 'BT': 71, 'BU': 72, 'BV': 73, 'BW': 74, 'BX': 75,
                     'BY': 76, 'BZ': 77, 'CA': 78, 'CB': 79, 'CC': 80, 'CD': 81, 'CE': 82,
                     'CF': 83, 'CG': 84, 'CH': 85, 'CI': 86, 'CJ': 87, 'CK': 88, 'CL': 89,
                     'CM': 90, 'CN': 91, 'CO': 92, 'CP': 93, 'CQ': 94, 'CR': 95, 'CS': 96,
                     'CT': 97, 'CU': 98, 'CV': 99, 'CW': 100, 'CX': 101, 'CY': 102, 'CZ': 103}
            return Excel[x]

        Target = Azbuka(self, App.get_running_app().root.ids.Target.text)  # Цель
        TInn = Azbuka(self, App.get_running_app().root.ids.TInn.text)  # ИНН цели
        Inn = Azbuka(self, App.get_running_app().root.ids.Inn.text)  # ИНН контрагента
        Partner = Azbuka(self, App.get_running_app().root.ids.Partner.text)  # Контрагент
        Corrbank = Azbuka(self, App.get_running_app().root.ids.Corrbank.text)  # Банк
        Account = Azbuka(self, App.get_running_app().root.ids.Account.text)  # Счет контрагента
        TAccount = Azbuka(self, App.get_running_app().root.ids.TAccount.text)  # Счет цели
        Debet = Azbuka(self, App.get_running_app().root.ids.Debet.text)  # Дебет
        Credit = Azbuka(self, App.get_running_app().root.ids.Credit.text)  # Кредит
        Data = Azbuka(self, App.get_running_app().root.ids.Data.text)  # Дата
        Description = Azbuka(self, App.get_running_app().root.ids.Description.text)  # Назначение
        files = self.file_paths
        table = pd.read_excel(files['Data'], header=None)
        try:
            inn_table = pd.read_excel(files['Inn'], header=None, skiprows=1)
        except:
            inn_table = None
        # убираем лишние \n
        table = table.replace(r'\n', '', regex=True)
        table = table.replace(r'-', '', regex=True)
        try:
            inn_table = inn_table.replace(r'\n', '', regex=True)
            inn_table = inn_table.replace(r' ', '', regex=True)
            inn_table[1] = 1
        except:
            None

        def stroka():
            i = 1
            while 1 == 1:
                # Возможно по счету не надо делать. Надо посмотреть на другие формы выписки
                if len(table.loc[i:i, Acc].replace(' ', '', regex=True).squeeze()) == 5:
                    return i
                i += 1
                if i == 50:
                    break

        # Считаем оборот
        Oborot = len(table.columns) + 1
        table[Debet] = table[Debet].replace(',', '.', regex=True)
        table[Credit] = table[Credit].replace(',', '.', regex=True)
        table[Debet] = pd.to_numeric(table[Debet], errors='coerce').fillna(0).astype(float)
        table[Credit] = pd.to_numeric(table[Credit], errors='coerce').fillna(0).astype(float)
        table[Oborot] = table[Debet] + table[Credit]
        # Platezh = 'Oborot+-'
        Platezh = Oborot + 1
        table[Platezh] = table[Credit] - table[Debet]
        # ЛЕВСИМВ к номеру счета
        table[Account] = pd.to_numeric(table[Account], errors='coerce').fillna(0).astype(str)
        Acc = len(table.columns) + 1
        table[Acc] = table[Account].str[:5]
        # Определяем какая строка является первой, содержащей данные (берем за основу столбик со счетами)
        Stroka = stroka()

        # Решаем проблему с DateStamp
        def DataStamp_error():
            Poterya = table[Data].iloc[0:Stroka].copy()
            try:
                table[Data] = table[Data][Stroka:].apply(lambda x: pd.to_datetime(x).strftime('%d.%m.%Y'))
                table[Data][0:Stroka] = Poterya
            except:
                pass

        DataStamp_error()
        # Проставляем заемщиков
        try:
            inn_table.columns = [Inn, 'loaner']
            inn_table[Inn] = pd.to_numeric(inn_table[Inn], errors='coerce').fillna(0).astype(float)
            table['ИНН+'] = table[Inn]
            table['ИНН+'] = pd.to_numeric(table['ИНН+'], errors='coerce').fillna(0).astype(float)
            table = pd.merge(table, inn_table, how='left', left_on='ИНН+', right_on=Inn, suffixes=('', '_y'))
        except:
            None
        # пример использования фильтра: graph_table.loc[lambda graph_table: graph_table['loaner'] == 'Заемщик', 'loaner'] = 'dict(x='+ graph_table['Dataplus'] + ", y=" + graph_table[Platezh].astype(str) + ", xref='x', yref='y', text='', showarrow=True, arrowhead=7, ax=0, ay=0)"
        try:
            table.drop('ИНН+', axis=1, inplace=True)
            table.drop(str(Inn) + '_y', axis=1, inplace=True)
            table.rename(columns={str(Inn): Inn}, inplace=True)
        except:
            None
        # отображаемый клиент
        nonTTInn = App.get_running_app().root.ids.targetInn.text
        try:
            if len(self.inns.columns) > 1:
                try:
                    TTInn = int(self.inns.loc[self.inns[1] == str(nonTTInn)][0])
                except:
                    TTInn = self.inns.loc[self.inns[1] == str(nonTTInn)][0]
            else:
                try:
                    TTInn = nonTTInn.astype(int)
                except:
                    TTInn = nonTTInn
        except:
            try:
                TTInn = nonTTInn.astype(int)
            except:
                TTInn = nonTTInn
        Ttable = table[table[TInn] == TTInn]
        import plotly.offline as py
        import plotly.figure_factory as ff
        import plotly.graph_objs as go
        try:
            graph_table = Ttable[[Target, TInn, Data, Partner, Debet, Credit, Oborot, Description, Acc, Platezh, Inn,
                                  'loaner']][Stroka - 1:]
        except:
            graph_table = Ttable[[Target, TInn, Data, Partner, Debet, Credit, Oborot, Description, Acc, Platezh, Inn]][
                          Stroka - 1:]
        graph_table.reset_index(inplace=True)
        graph_table['nuli'] = 0

        def data_no_repeat():
            graph_table['Dataplus'] = str(graph_table[Data].replace(' ', '', regex=True))
            ii = 1
            for i in range(graph_table.shape[0]):
                if i == 0:
                    graph_table.loc[i:i, 'Dataplus'] = str(
                        graph_table.loc[i:i, Data].replace(' ', '', regex=True).squeeze()).replace(' 00:00:00',
                                                                                                   '').replace('-', '.')
                    continue
                if graph_table.loc[i - 1:i - 1, Data].squeeze() == graph_table.loc[i:i, Data].squeeze():
                    graph_table.loc[i:i, 'Dataplus'] = str(
                        graph_table.loc[i:i, Data].replace(' ', '', regex=True).squeeze()).replace(' 00:00:00',
                                                                                                   '').replace('-',
                                                                                                               '.') + "_" + str(
                        ii)
                    ii = ii + 1
                else:
                    graph_table.loc[i:i, 'Dataplus'] = str(
                        graph_table.loc[i:i, Data].replace(' ', '', regex=True).squeeze()).replace(' 00:00:00',
                                                                                                   '').replace('-', '.')
                    ii = 1

        data_no_repeat()
        graph_table['color'] = 'rgb(89, 96, 109)'
        # свод по крупнейшим контрагентам
        pivot = Ttable.pivot_table([Debet, Credit, Oborot, Account], index=[Partner],
                                   aggfunc={Debet: 'sum', Credit: 'sum', Oborot: 'sum', Account: 'count'},
                                   fill_value=0).sort_values(by=Oborot, ascending=False)
        pivot.columns = ['ДО', 'КО', 'Число операций', 'Оборот']
        Bigpartners = pivot.iloc[0:10, 0:0]
        Bigpartners['color'] = ['rgb(0, 0, 50)', 'rgb(0, 0, 100)', 'rgb(0, 0, 150)', 'rgb(0, 0, 200)',
                                'rgb(0, 50, 200)', 'rgb(0, 100, 200)', 'rgb(0, 150, 200)', 'rgb(0, 200, 200)',
                                'rgb(0, 200, 250)', 'rgb(0, 250, 250)']
        Bigpartners.reset_index(inplace=True)
        graph_table = pd.merge(graph_table, Bigpartners, how='left', left_on=Partner, right_on=Partner)

        def kraska():
            # находим крупнейших контрагентов и даем им оттенки синего
            for i in range(graph_table.shape[0]):
                if graph_table.loc[i:i, Acc].astype(str).str[:2].squeeze() == '45':
                    graph_table.loc[i:i, 'color_y'] = 'rgb(0, 128, 0)'  # зеленый
                if graph_table.loc[i:i, 'color_y'].isnull().squeeze():
                    if graph_table.loc[i:i, Acc].squeeze() == '20202':
                        graph_table.loc[i:i, 'color_y'] = 'rgb(232, 205, 0)'  # коричневый
                    if graph_table.loc[i:i, Acc].squeeze() == '47427':
                        graph_table.loc[i:i, 'color_y'] = 'rgb(0, 128, 0)'  # зеленый
                    if graph_table.loc[i:i, Acc].squeeze() == '40817':
                        graph_table.loc[i:i, 'color_y'] = 'rgb(128, 83, 0)'  # коричневый
                    else:
                        graph_table.loc[i:i, 'color_y'] = graph_table.loc[i:i, 'color_x']

        kraska()
        diagramma = [go.Bar(x=graph_table.Dataplus, y=graph_table[Platezh], text=graph_table[Partner],
                            marker=dict(color=graph_table.color_y), name='Платеж', ),
                     go.Bar(x=graph_table.Dataplus, y=graph_table.nuli, text=graph_table[Description],
                            marker=dict(color='rgb(255, 255, 255)'), name='Назначение')]

        def annotation():
            annotations = []
            loans = graph_table[graph_table['loaner'].notnull()][['Dataplus', Platezh]]
            for i in loans.index:
                annotations.append(
                    dict(x=loans.loc[i:i, 'Dataplus'].squeeze(), y=loans.loc[i:i, Platezh].squeeze(), xref='x',
                         yref='y', text='', showarrow=True, arrowhead=7, ax=0, ay=0))
            return annotations

        layout = go.Layout(showlegend=True, annotations=annotation())
        fig = go.Figure(data=diagramma, layout=layout)
        py.plot(fig, config={'scrollZoom': True}, filename='~/Desktop/' + str(str(TTInn) + '.html'))

    # ОСНОВНАЯ ФУНКЦИЯ СБИВАТОРА
    def change_Vid(self, err=0, reset=''):
        def VPR_loop(df, acc=90, col1=1, col2=2,
                     nachalo=None):  # ОСНОВНОЙ ЦИКЛ: ИЩЕМ ПОКА НЕ НАЙДЕМ, ПЕРЕБИРАЯ ТОЧНОСТЬ (90%,80%, ... , N пользователя с точностью до 10%)
            def VPR(what, where, acc=90, col1=1, col2=2):  # ВПР одного значения по столбцу
                acc = acc / 100
                try:
                    try:
                        what = what.lower()
                    except:
                        None
                    for i, r in where.dropna(subset=[col1 - 1]).iterrows():
                        try:
                            wh = r[where.columns.tolist()[0]].lower()
                        except:
                            wh = r[where.columns.tolist()[0]]
                        if difflib.get_close_matches(what, [wh], cutoff=acc) != []:
                            # print(what, '->', df3.loc[i:i, col1-1])
                            return i
                except Exception as e:
                    print(e)
                    return None

            def VPRing(df, acc=90, col1=1, col2=2):  # Запуск ВПР по всем искомым значениям
                df_1 = df[['0_x', 'Расшифровка']]
                df_2 = df[[col1 - 1,
                           col2 - 1]]  # 2 колонка передается для того, чтобы работал iterrows в VPR (он с series не работает)
                for i, val in df_1.loc[df_1['Расшифровка'] == 'Не найдено'].iterrows():
                    j = val[0]
                    try:
                        vpr_row = VPR(j, df_2, acc, col1, col2)
                        if vpr_row is not None:
                            df.loc[i:i, col2 - 1] = df.loc[vpr_row:vpr_row, col2 - 1].squeeze()
                            df.loc[i:i, col1 - 1] = df.loc[vpr_row:vpr_row, col1 - 1].squeeze()
                            df.loc[i:i, 'Расшифровка'] = 'Найдено ~ на ' + str(round(acc)) + '%'
                    except Exception as e:
                        print(e)
                        None

            # print('vpr poshel')
            # сделать цикл запусков до целевой точности
            accs = [90, 80, 70, 60, 50, 40, 30, 20, 10]
            accs = [s for s in accs if s >= acc]  # acc_nearest = min(l,key=lambda x: abs(x-acc)) # найти ближайшее
            for i in accs:
                if 'Не найдено' in df.loc[:, 'Расшифровка'].tolist():
                    VPRing(df, i, col1, col2)
            # сохраняем либо в ексель, либо в csv (если больше 1 млн ячеек)
            konec = (dt.datetime.now() - nachalo).total_seconds()
            from os.path import expanduser as osexp
            try:
                writer = pd.ExcelWriter(osexp(r'~/Desktop/' + 'DjinnsGift_' + str(round(konec, 2)) + '.xlsx'), \
                                        date_format='DD.MM.YY', datetime_format='DD.MM.YY')
                workbook = writer.book
                workbook.formats[0].set_font_size(8)
                workbook.formats[0].set_font_name('Arial Narrow')
                workbook.formats[0].set_num_format('#,')
                df.to_excel(writer, sheet_name='1', index=False)
                writer.save()
            except Exception as e:
                print(e)
                df.to_csv(osexp(r'~/Desktop/' + 'DjinnsGift_' + str(round(konec, 2)) + '.csv'), sep=";", decimal=',',
                          index=False, encoding='utf-8-sig')

        def hint_show(kadr):
            iii = [App.get_running_app().root.ids.q1, App.get_running_app().root.ids.hint1, \
                   App.get_running_app().root.ids.hint2, App.get_running_app().root.ids.hint3]
            for i in iii:
                # ПРОЗРАЧНОСТЬ
                if kadr > len(iii):
                    i.opacity = 0
                    i.disabled = True
                    i.size_hint = (0, 0)
                    i.pos_hint = {'x': 1, 'y': 1}
                else:
                    if i == iii[kadr]:
                        i.opacity = 1
                        i.disabled = False
                        # РАСПОЛОЖЕНИЕ
                        if kadr == 0:
                            i.size_hint = (None, None)
                            i.pos_hint = {'x': 0.584, 'y': 0.39}
                            i.height = dp(150)
                            i.width = dp(100)
                        if kadr == 1:
                            i.size_hint = (0.2, None)
                            i.pos_hint = {'x': 0.56, 'y': 0.2}
                            i.height = dp(90)
                        if kadr == 2:
                            i.size_hint = (0.3, None)
                            i.pos_hint = {'x': 0.59, 'y': 0.45}
                            i.height = dp(90)
                        if kadr == 3:
                            i.size_hint = (0.2, None)
                            i.pos_hint = {'x': 0.66, 'y': 0.4}
                            i.height = dp(90)

        def next_video(n=0):
            # СМЕНА ВИДЕО В СООТВЕТСТВИИ С КАДРОМ (STATE)
            App.get_running_app().root.ids.video.vid = SmartPlot.vid_list[SmartPlot.vid_state - n]
            App.get_running_app().root.ids.video.state = 'play'

        if reset == 1:
            hint_show(9)
            SmartPlot.vid_state = 1
            next_video(1)
            return None

        # ПРОВЕРКА НА ПРОИГРЫВАНИЕ ВИДЕО, если оно не проиграно полностью - перемотать и показать надпись папируса мгновенно
        if App.get_running_app().root.ids.video.position / App.get_running_app().root.ids.video.duration < 0.95 and SmartPlot.vid_state != 1:
            if SmartPlot.vid_state == 2:
                print('stop ' + str(SmartPlot.vid_state))
                App.get_running_app().root.ids.video.seek(
                    1)  # ЕСЛИ ВИДЕО ЕЩЕ ПРОИГРЫВАЕТСЯ -- ПЕРЕМОТАТЬ ДО КОНЦА (ЧТОБЫ НЕ ЖДАТЬ)
                SmartPlot.event.cancel()  # прервать отложенный показ текста к папирусу
                Clock.schedule_once(lambda x: hint_show(0), 0.4)

        else:
            # ЕСЛИ В ПРОГРАММЕ НЕТ ОШИБКИ, ТО КАДРЫ С ОШИБКОЙ ПРОПУСКАЮТСЯ
            if err == 0:
                if SmartPlot.vid_state == 3 or SmartPlot.vid_state == 5 or SmartPlot.vid_state == 7:
                    SmartPlot.vid_state += 1
            print(SmartPlot.vid_state)

            # ПОВЕДЕНИЕ НА ТОМ ИЛИ ИНОМ КАДРЕ (STATE)
            if SmartPlot.vid_state == 1:  # экран лампы
                hint_show(9)  # ВСЕ НАДПИСИ KIVY СКРЫВАЕМ
                SmartPlot.event = Clock.schedule_once(lambda x: hint_show(0), 9.5)
                SmartPlot.event()
                next_video()
            if SmartPlot.vid_state == 2:  # экран появления
                hint_show(9)  # ВСЕ НАДПИСИ KIVY СКРЫВАЕМ
                SmartPlot.event = Clock.schedule_once(lambda x: hint_show(1), 3)
                SmartPlot.event()
                next_video()
            if SmartPlot.vid_state == 4:  # экран левой руки
                try:
                    if SmartPlot.df1 is None:
                        SmartPlot.df1 = pd.read_clipboard('\t', header=None)
                        print(SmartPlot.df1)
                    elif len(SmartPlot.df1.columns) > 1 or SmartPlot.df1 is None:
                        Error_sim = 5 / 0  # симуляция ошибки
                    else:
                        print(SmartPlot.df1)
                    next_video()
                    hint_show(9)  # ВСЕ НАДПИСИ KIVY СКРЫВАЕМ
                    SmartPlot.event = Clock.schedule_once(lambda x: hint_show(2), 1)
                    SmartPlot.event()
                except Exception as e:
                    print('ОШИБКА!:')
                    print(e)
                    next_video(1)
                    SmartPlot.vid_state -= 1
                    SmartPlot.event = Clock.schedule_once(lambda x: hint_show(1), 2)
                    SmartPlot.event()
            if SmartPlot.vid_state == 6:  # экран правой руки ##########################################################################################################################################
                try:
                    if SmartPlot.df2 is None:
                        SmartPlot.df2 = pd.read_clipboard('\t', header=None)
                        print(SmartPlot.df2)
                        if len(SmartPlot.df2.columns) < 2 or SmartPlot.df2 is None:
                            Error_sim = 5 / 0  # симуляция ошибки
                    else:
                        print(SmartPlot.df2)

                    # В НАЧАЛЕ ДЕЛАЕМ ТОЧНЫЙ ПОИСК
                    col1 = int(App.get_running_app().root.ids.vpr_col1.text)  # 3 # где ищем
                    col2 = int(App.get_running_app().root.ids.vpr_col2.text)  # 5 # что тянем
                    acc = int(App.get_running_app().root.ids.vpr_acc.text)  # точность поиска
                    nachalo = dt.datetime.now()  # засекаем время начала просчетов
                    df3 = pd.merge(SmartPlot.df1, SmartPlot.df2, how='outer', left_on=0,
                                   right_on=col1 - 1)  # объединяем со 100% совпадением
                    df3 = df3.dropna(how='all')  # убираем полностью пустые столбцы и строки
                    # РАСШИФРОВЫВАЕМ РЕЗУЛЬТАТЫ -- ПРОСТАВЛЯЕМ МАРКЕР
                    try:
                        df3.loc[df3['0_x'].isna(), 'Расшифровка'] = 'Лишнее'
                        df3.loc[
                            (df3['0_x'] != df3[col1 - 1]) & (df3['Расшифровка'].isna()), 'Расшифровка'] = 'Не найдено'
                        df3['id'] = pd.factorize(df3.loc[:, '0_x'])[0]
                        df3.loc[
                            (df3['0_x'].duplicated()) & (df3['Расшифровка'].isna()), 'Расшифровка'] = 'Несколько 100%'
                        df3.loc[
                            (df3['0_x'] == df3[col1 - 1]) & (df3['Расшифровка'].isna()), 'Расшифровка'] = 'Найдено 100%'
                    except:
                        df3.loc[df3[0].isna(), 'Расшифровка'] = 'Лишнее'
                        df3.loc[
                            (df3[0] != df3[col1 - 1]) & (df3['Расшифровка'].isna()), 'Расшифровка'] = 'Не найдено'
                        df3['id'] = pd.factorize(df3.loc[:, 0])[0]
                        df3.loc[
                            (df3[0].duplicated()) & (df3['Расшифровка'].isna()), 'Расшифровка'] = 'Несколько 100%'
                        df3.loc[
                            (df3[0] == df3[col1 - 1]) & (df3['Расшифровка'].isna()), 'Расшифровка'] = 'Найдено 100%'
                    # ЕСЛИ НУЖНО ДЕЛАЕМ НЕТОЧНЫЙ ПОИСК
                    if acc < 100:
                        try:
                            VPR_loop(df3, acc, col1, col2, nachalo)  # желательно на другом ядре
                        except Exception as e:
                            print(e)
                            try:
                                from os.path import expanduser as osexp
                                writer = pd.ExcelWriter(osexp(r'~/Desktop/' + 'DjinnsGift' + '.xlsx'), \
                                                        date_format='DD.MM.YY', datetime_format='DD.MM.YY')
                                workbook = writer.book
                                workbook.formats[0].set_font_size(8)
                                workbook.formats[0].set_font_name('Arial Narrow')
                                workbook.formats[0].set_num_format('#,')
                                df3.to_excel(writer, sheet_name='1', index=False)
                                writer.save()
                            except Exception as e:
                                print(e)
                                df3.to_csv(osexp(r'~/Desktop/' + 'DjinnsGift' + '.csv'), sep=";", decimal=',',
                                           index=False, encoding='utf-8-sig')
                    else:  # Если примерный ВПР не нужен, сохраняем точные результаты
                        try:
                            from os.path import expanduser as osexp
                            writer = pd.ExcelWriter(osexp(r'~/Desktop/' + 'DjinnsGift' + '.xlsx'), \
                                                    date_format='DD.MM.YY', datetime_format='DD.MM.YY')
                            workbook = writer.book
                            workbook.formats[0].set_font_size(8)
                            workbook.formats[0].set_font_name('Arial Narrow')
                            workbook.formats[0].set_num_format('#,')
                            df3.to_excel(writer, sheet_name='1', index=False)
                            writer.save()
                        except Exception as e:
                            print(e)
                            df3.to_csv(osexp(r'~/Desktop/' + 'DjinnsGift' + '.csv'), sep=";", decimal=',', index=False,
                                       encoding='utf-8-sig')
                    # e = 5/0# симуляция ошибки
                    next_video()
                    hint_show(9)  # ВСЕ НАДПИСИ KIVY СКРЫВАЕМ
                    SmartPlot.event = Clock.schedule_once(lambda x: hint_show(3), 1)
                    SmartPlot.event()
                except Exception as e:
                    print('ОШИБКА!:')
                    print(e)
                    from os.path import expanduser as osexp
                    df3.to_csv(osexp(r'~/Desktop/' + 'DjinnsGift_err' + '.csv'), sep=";", decimal=',', index=False,
                               encoding='utf-8-sig')
                    next_video(1)
                    SmartPlot.vid_state -= 1
                    SmartPlot.event = Clock.schedule_once(lambda x: hint_show(2), 1.5)
                    SmartPlot.event()
            if SmartPlot.vid_state == 8:  # Экран кристалла
                try:
                    # e = 5/0# симуляция ошибки
                    next_video()
                    hint_show(9)  # ВСЕ НАДПИСИ KIVY СКРЫВАЕМ
                except Exception as e:
                    print('ОШИБКА!:')
                    print(e)
                    next_video(1)
                    hint_show(9)  # ВСЕ НАДПИСИ KIVY СКРЫВАЕМ
                    SmartPlot.vid_state == 10  # НИЧЕГО НЕ ПОЛУЧИЛОСЬ
            if SmartPlot.vid_state == 10:  # ЕСЛИ НИЧЕГО НЕ ПОЛУЧИЛОСЬ
                SmartPlot.vid_state = 0
            else:
                hint_show(9)  # ВСЕ НАДПИСИ KIVY СКРЫВАЕМ

            if SmartPlot.vid_state == len(
                    SmartPlot.vid_list) - 1:  # чтобы после окончания последнего видео всё начиналось сначала
                SmartPlot.vid_state = 1
            else:
                SmartPlot.vid_state += 1

    # ОСНОВНАЯ ФУНКЦИЯ ТИПОГРАФА
    def Typograph(self):
        if App.get_running_app().root.ids.rubl.state == 'down':
            need_rub = True
        else:
            need_rub = False
        if App.get_running_app().root.ids.ndate.state == 'down':
            need_date = True
        else:
            need_date = False
        try:
            path = r'' + SmartPlot.file_paths.decode("utf-8")
            import docx
            import re
            def data_ch(txt):
                txt = txt.group()
                mounths = {1: 'января', 2: 'февраля', 3: 'марта', 4: 'апреля', 5: 'мая', 6: 'июня', 7: 'июля',
                           8: 'августа',
                           9: 'сентября', 10: 'октября', 11: 'ноября', 12: 'декабря'}
                try:
                    data_f = re.findall(r'\d\d\.\d\d\.\d\d\d\d', txt)[0].split('.')
                    data_f_adj = str(int(data_f[0])) + chr(160) + mounths[int(data_f[1])] + chr(160) + str(
                        int(data_f[2])) + chr(
                        160) + 'г.'
                    return data_f_adj
                except Exception as e:
                    print(e)
                    try:
                        data_f = re.findall(r'\d\.\d\d\.\d\d\d\d', txt)[0].split('.')
                        data_f_adj = ' ' + str(int(data_f[0])) + chr(160) + mounths[int(data_f[1])] + chr(160) + str(
                            int(data_f[2])) + chr(
                            160) + 'г.'
                        return data_f_adj
                    except Exception as e:
                        print(e)
                        try:
                            data_f = re.findall(r'\d\.\d\.\d\d\d\d', txt)[0].split('.')
                            data_f_adj = ' ' + str(int(data_f[0])) + chr(160) + mounths[int(data_f[1])] + chr(
                                160) + str(
                                int(data_f[2])) + chr(
                                160) + 'г.'
                            return data_f_adj
                        except Exception as e:
                            print(e)
                            try:
                                data_f = re.findall(r'\d\d\.\d\.\d\d\d\d', txt)[0].split('.')
                                data_f_adj = str(int(data_f[0])) + chr(160) + mounths[int(data_f[1])] + chr(160) + str(
                                    int(data_f[2])) + chr(
                                    160) + 'г.'
                                return data_f_adj
                            except Exception as e:
                                print(e)
                                return str(txt)

            def kavicho_ch(txt):  # open
                txt = txt.group()
                return '«' + re.findall(r'"\w', txt)[0][1]

            def kavichc_ch(txt):  # close
                txt = txt.group()
                return '»' + re.findall(r'"\W', txt)[0][1]

            def sps_ch(txt):  # close
                txt = txt.group()
                return re.findall(r'\d ', txt)[0][0] + chr(160)  # ' '

            def typograph(i, ii):
                ### КАВЫЧКИ
                if '"' in i.text:
                    i.text = re.sub(r'"\w', kavicho_ch, i.text)
                    i.text = re.sub(r'"\W', kavichc_ch, i.text)
                if '"' in i.text:  # если p.runs разделил " и следующий символ
                    i_txt = i.text + p.runs[ii + 1].text[0]
                    i.text = re.sub(r'"\w', '«', i_txt)  # [:-1]
                    i_txt = i.text + p.runs[ii + 1].text[0]
                    i.text = re.sub(r'"\W', '»', i_txt)[:-1]
                ### ТИРЕ ВМЕСТО ДЕФИСА
                if '-' in i.text:
                    i.text = i.text.replace(' - ', chr(160) + '– ')  # '—' + '–'
                ### ЧАСТЫЕ ОШИБКИ
                i.text = i.text.replace('млн.', 'млн')
                i.text = i.text.replace('млнруб', 'млн' + chr(160) + 'руб')
                i.text = i.text.replace('млрд.', 'млрд')
                i.text = i.text.replace('млрдруб', 'млрд' + chr(160) + 'руб')
                i.text = i.text.replace('тыс ', 'тыс.' + chr(160))
                i.text = i.text.replace('млрдруб', 'млрд' + chr(160) + 'руб')
                ### ДВОЙНЫЕ ПРОБЕЛЫ
                ### НЕРАЗРЫВНЫЕ ПРОБЕЛЫ
                i.text = i.text.replace('  ', ' ')
                if re.findall(r'\d ', i.text):
                    i.text = re.sub(r'\d ', sps_ch, i.text)
                i.text = i.text.replace('млн ', 'млн' + chr(160))
                i.text = i.text.replace('млрд ', 'млрд' + chr(160))
                i.text = i.text.replace('тыс. ', 'тыс.' + chr(160))
                i.text = i.text.replace('ООО ', 'ООО' + chr(160))
                i.text = i.text.replace('ЗАО ', 'ЗАО' + chr(160))
                i.text = i.text.replace('НАО ', 'НАО' + chr(160))
                i.text = i.text.replace('ОАО ', 'ОАО' + chr(160))
                i.text = i.text.replace('ПАО ', 'ПАО' + chr(160))
                i.text = i.text.replace('АО ', 'АО' + chr(160))

                if need_date:
                    ### ДАТЫ
                    # print('Обработка дат')
                    if re.findall(r'\d\.\d\d\d\d', i.text):
                        i.text = re.sub(r'\d\d\.\d\d\.\d\d\d\d', data_ch, i.text)
                        i.text = re.sub(r'\d\d\.\d\.\d\d\d\d', data_ch, i.text)
                        i.text = re.sub(r' \d\.\d\d\.\d\d\d\d', data_ch, i.text)
                        i.text = re.sub(r' \d\.\d\.\d\d\d\d', data_ch, i.text)
                        i.text = i.text.replace('..', '.')
                        i.text = i.text.replace(' .', '.')
                        i.text = i.text.replace(chr(160) + '.', '.')
                if need_rub:
                    i.text = i.text.replace('руб.', '₽')
                if len(p.text) > 1 and p.text[-1] == '₽':
                    p.text = p.text + '.'

            doc = docx.Document(path)
            for p in doc.paragraphs:
                ii = 0
                if 'graphicData' in p._p.xml:
                    continue
                for i in p.runs:
                    if i.text != '':
                        typograph(i, ii)
                    ii += 1
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            ii = 0
                            if 'graphicData' in p._p.xml:
                                continue
                            for i in p.runs:
                                if i.text != '':
                                    typograph(i, ii)
                                ii += 1
            # СОХРАНЕНИЕ
            from os.path import expanduser as osexp
            try:
                doc.save(osexp(r'~/Desktop/' + 'Report' + '.docx'))
            except:
                doc.save(osexp(r'~/Рабочий стол/' + 'Report' + '.docx'))
        except Exception as e:
            print(e)
            Popup().open()
        print('Done')

    # ОСНОВНАЯ ФУНКЦИЯ НЕДВИГЕРА
    def EGRN(self):
        def parse(path):
            try:
                tree = etree.parse(path)
                lstKey = []
                lstValue = []
                for p in tree.iter():
                    lstKey.append(tree.getpath(p).replace("/", ".")[1:])
                    lstValue.append(p.text)
                df = pd.DataFrame({'key': lstKey, 'value': lstValue})
                df.loc[:, 'Кадастровый№'] = df.loc[df['key'].str.contains('CadastralNumber'), 'value'].squeeze()
                df.loc[:, 'Тип'] = df.loc[df['key'].str.contains('ObjectDesc.Name'), 'value'].squeeze()
                df.loc[:, 'Подтип'] = df.loc[df['key'].str.contains('ObjectDesc.ObjectTypeText'), 'value'].squeeze()
                df.loc[:, 'Категория'] = df.loc[
                    df['key'].str.contains('ObjectDesc.GroundCategoryText'), 'value'].squeeze()
                df.loc[:, 'Площадь'] = df.loc[df['key'].str.contains('ObjectDesc.Area.AreaText'), 'value'].squeeze()
                df.loc[:, 'Адрес'] = df.loc[df['key'].str.contains('ObjectDesc.Address.Content'), 'value'].squeeze()
                # Заполняем собственника
                df.loc[df['key'].str.contains('Owner.*Content', regex=True), 'Собственник'] = df.loc[
                    df['key'].str.contains('Owner.*Content', regex=True), 'value']
                df.loc[:, 'Собственник'].fillna(method='ffill', inplace=True)
                # заполняем право
                df.loc[df['key'].str.contains('Registration.*Name', regex=True) & ~df['key'].str.contains('DocFound',
                                                                                                          regex=True), 'Право'] = \
                df.loc[df['key'].str.contains('Registration.*Name', regex=True) & ~df['key'].str.contains('DocFound',
                                                                                                          regex=True), 'value']
                # заполняем дату регистрации
                df.loc[df['key'].str.contains('Registration.*RegDate', regex=True), 'Дата регистрации права'] = df.loc[
                    df['key'].str.contains('Registration.*RegDate', regex=True), 'value']
                # заполняем № регистрации
                df.loc[df['key'].str.contains('Registration.*RegNumber', regex=True), '№ регистрации права'] = df.loc[
                    df['key'].str.contains('Registration.*RegNumber', regex=True), 'value']
                # заполняем Правоустанавливающий документ
                df.loc[df['key'].str.contains('Registration.*Name', regex=True) & df['key'].str.contains('DocFound',
                                                                                                         regex=True), 'Правоустанавливающий документ'] = \
                df.loc[df['key'].str.contains('Registration.*Name', regex=True) & df['key'].str.contains('DocFound',
                                                                                                         regex=True), 'value']
                # заполняем № правоустанавливающего документа
                df.loc[df['key'].str.contains('Registration.*Number', regex=True) & df['key'].str.contains('DocFound',
                                                                                                           regex=True), '№ правоустанавливающего документа'] = \
                df.loc[df['key'].str.contains('Registration.*Number', regex=True) & df['key'].str.contains('DocFound',
                                                                                                           regex=True), 'value']
                # заполняем Дата правоустанавливающего документа
                df.loc[df['key'].str.contains('Registration.*Date', regex=True) & df['key'].str.contains('DocFound',
                                                                                                         regex=True), 'Дата правоустанавливающего документа'] = \
                df.loc[df['key'].str.contains('Registration.*Date', regex=True) & df['key'].str.contains('DocFound',
                                                                                                         regex=True), 'value']
                # заполняем NULL
                owners = df.loc[:, 'Собственник'].dropna().unique().tolist()
                for owner in owners:
                    for column in df.columns.tolist()[9:]:
                        df.loc[df['Собственник'] == owner, column] = df.loc[
                            (df['Собственник'] == owner) & (df[column].isna() == False), column].squeeze()
                # Фильтруем, убираем лишние столбцы
                df = df.loc[df['key'].str.contains('Owner.*Content', regex=True), df.columns.tolist()[2:]]
                return df
            except:
                print('TROUBLE!!!!!!!!!!!!!', path)
        def parsing(fpath):
            files = gb(fpath)
            res = []
            for i, file in enumerate(files):
                print(file)
                if i == 0:
                    result = parse(file)
                else:
                    result = pd.concat([result, parse(file)])
                result.to_excel('!KAD.xls', index=False)

        try:
            path = r'' + SmartPlot.file_paths.decode("utf-8") + r"\*.xml"
            parsing(path)
            print('Сохранено!')
        except Exception as e:
            print(e)
            Popup().open()


class DropBut(Button):
    def __init__(self, **kwargs):
        super(DropBut, self).__init__(**kwargs)
        self.drop_list = None
        self.drop_list = DropDown()

        def upd_droplist(self):
            self.drop_list.clear_widgets()
            if App.get_running_app().root.current_tab.text == "Выпискатор":
                types = SmartPlot.inn_table
            else:
                types = ['в рублях', 'в тысячах', 'в миллионах', 'в миллиардах']
            if types is not None:
                for i in types:
                    btn = Button(text=str(i), size_hint_y=None, height=30, font_size=12,
                                 text_size=(self.width - 10, None), halign='left', valign='middle',
                                 color=[0.3, 0.3, 0.3, 1], background_color=[256, 256, 256, 0.7])
                    btn.bind(on_release=lambda btn: self.drop_list.select(btn.text))
                    self.drop_list.add_widget(btn)

        self.bind(on_release=upd_droplist)
        self.bind(on_release=self.drop_list.open)

        def vibrano(item):
            if App.get_running_app().root.current_tab.text == "Выпискатор":
                App.get_running_app().root.ids.targetInn.text = item
                print(inn_table)
            else:
                App.get_running_app().root.ids.izm.text = item

        self.drop_list.bind(on_select=lambda instance, x: vibrano(x))


if __name__ == '__main__':
    SmartPlot().run()
